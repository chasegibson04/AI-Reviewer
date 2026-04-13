#!/usr/bin/env python3
from __future__ import annotations

import argparse
import json
import re
import shutil
from dataclasses import dataclass
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from ai_reviewer.eval.comment_quality import compare_metric_dicts, compute_comment_quality_metrics
from ai_reviewer.ingest.types import ParsedDocument
from ai_reviewer.models.ollama_provider import OllamaProvider
from ai_reviewer.review.comment_responder import (
    render_comment_response_manifest_markdown,
    run_existing_comment_responder,
)
from ai_reviewer.review.manuscript_annotation import _build_section_index_map, _analysis_paragraphs, build_annotated_manuscript_output
from ai_reviewer.review.sentence_claim_check import MockExternalSearchProvider, run_sentence_claim_check
from ai_reviewer.tools.docx_tools import get_docx_paragraph_texts
from docx import Document


ALLOWED_CASES = [
    {
        "project_id": "20260325163524_test-existingphactorpaper",
        "source_docx": "projects/20260325163524_test-existingphactorpaper/materials/managed/20260329_174704_137524/project1_clean_native.docx",
        "review_json": "projects/20260325163524_test-existingphactorpaper/runs/20260330_224209_deep_run/final_deep_review_report.json",
    },
    {
        "project_id": "20260327051312_miniaturization_d2b",
        "source_docx": "projects/20260327051312_miniaturization_d2b/materials/managed/20260329_174704_158099/project2_clean_native.docx",
        "review_json": "projects/20260327051312_miniaturization_d2b/runs/20260329_122934_deep_run/final_deep_review_report.json",
    },
]

COMPARISON_KEYS = [
    "anchor_localization.avg_anchor_length_words",
    "anchor_localization.max_anchor_length_words",
    "anchor_localization.sentence_local_exceed_one_sentence_rate",
    "anchor_localization.paragraph_wide_local_anchor_rate",
    "anchor_localization.clause_local_too_broad_rate",
    "rewrite_usefulness.no_op_rewrite_rate",
    "rewrite_usefulness.trivial_rewrite_rate",
    "rewrite_usefulness.style_suppressed_rewrite_rate",
    "rewrite_usefulness.materially_better_surviving_rate",
    "comment_concision.avg_comment_length_words",
    "comment_concision.comments_above_threshold_rate",
    "comment_concision.actionable_to_generic_ratio",
    "deduplication_quality.near_duplicate_pair_rate",
    "deduplication_quality.same_sentence_overlap_count",
    "deduplication_quality.overlapping_comment_rate",
    "deduplication_quality.merged_total",
    "deduplication_quality.unmerged_duplicate_pairs_final",
    "claim_check_coverage.sentences_classified",
    "claim_check_coverage.factual_claim_count",
    "claim_check_coverage.citation_needed_rate",
    "claim_check_coverage.overclaim_rate",
    "claim_check_coverage.unclear_support_rate",
    "claim_check_coverage.surfaced_comment_count",
    "claim_check_coverage.suppressed_comment_count",
    "claim_check_coverage.surfaced_high_value_rate",
    "style_pass_outcomes.style_comments_surfaced_count",
    "style_pass_outcomes.style_rewrites_accepted_count",
    "style_pass_outcomes.style_rewrites_suppressed_count",
    "privacy_search_safety.safe_externalization_count",
    "privacy_search_safety.blocked_externalization_count",
    "existing_comment_responder_coverage.source_comments_read_count",
    "existing_comment_responder_coverage.effective_source_comments_read_count",
    "existing_comment_responder_coverage.concrete_response_count",
    "existing_comment_responder_coverage.needs_clarification_count",
    "overall_comment_mix.rates.style/clarity",
    "overall_comment_mix.rates.evidence/overclaim",
    "overall_comment_mix.rates.wording/precision",
    "overall_comment_mix.rates.grammar/mechanics",
    "overall_comment_mix.rates.structure/flow",
]


@dataclass
class ReviewObj:
    payload: dict[str, Any]

    @property
    def section_specific_comments(self):
        return [
            type(
                "SectionComment",
                (),
                {
                    "section": row.get("section", "section"),
                    "comment": row.get("comment", ""),
                    "severity": row.get("severity", "medium"),
                    "evidence_source": row.get("evidence_source"),
                    "manuscript_quote": row.get("manuscript_quote"),
                },
            )
            for row in self.payload.get("section_specific_comments", [])
            if isinstance(row, dict)
        ]

    @property
    def extracted_action_items(self):
        return [
            type("ActionItem", (), {"action": row.get("action", ""), "priority": row.get("priority", "medium")})
            for row in self.payload.get("extracted_action_items", [])
            if isinstance(row, dict)
        ]

    @property
    def detailed_reviewer_comments(self):
        return [str(x) for x in self.payload.get("detailed_reviewer_comments", []) if str(x).strip()]

    @property
    def grounded_detailed_comments(self):
        return [
            type(
                "GroundedComment",
                (),
                {
                    "comment": row.get("comment", ""),
                    "severity": row.get("severity", "medium"),
                    "evidence_source": row.get("evidence_source"),
                    "manuscript_quote": row.get("manuscript_quote"),
                },
            )
            for row in self.payload.get("grounded_detailed_comments", [])
            if isinstance(row, dict)
        ]

    @property
    def methodological_concerns(self):
        return [str(x) for x in self.payload.get("methodological_concerns", []) if str(x).strip()]

    @property
    def novelty_concerns(self):
        return [str(x) for x in self.payload.get("novelty_concerns", []) if str(x).strip()]

    @property
    def citation_reference_concerns(self):
        return [str(x) for x in self.payload.get("citation_reference_concerns", []) if str(x).strip()]

    @property
    def writing_organization_concerns(self):
        return [str(x) for x in self.payload.get("writing_organization_concerns", []) if str(x).strip()]

    @property
    def figure_table_concerns(self):
        return [str(x) for x in self.payload.get("figure_table_concerns", []) if str(x).strip()]

    @property
    def major_weaknesses(self):
        return [str(x) for x in self.payload.get("major_weaknesses", []) if str(x).strip()]


def _build_doc(base_docx: Path) -> ParsedDocument:
    paragraphs = get_docx_paragraph_texts(base_docx, normalize_review_artifacts=True)
    text = "\n\n".join(p.strip() for p in paragraphs if p.strip())
    now = datetime.now(timezone.utc).isoformat()
    return ParsedDocument(
        source_path=base_docx,
        source_path_abs=base_docx.resolve(),
        source_path_rel=str(base_docx),
        file_size_bytes=base_docx.stat().st_size,
        document_type="docx",
        parse_engine="comment_quality_eval",
        ingest_timestamp=now,
        raw_text=text,
        cleaned_text=text,
        headings=[],
        page_count=None,
        parse_warnings=[],
        chunks=[],
    )


def _run_privacy_probe(base_docx: Path) -> dict[str, Any]:
    paragraphs = _analysis_paragraphs(base_docx)
    section_by_idx = _build_section_index_map(paragraphs)
    mock_provider = MockExternalSearchProvider({})
    manifest = run_sentence_claim_check(
        paragraphs=paragraphs,
        section_by_idx=section_by_idx,
        provider=None,
        model=None,
        timeout_seconds=60,
        supporting_cards=None,
        external_search_enabled=True,
        external_provider=mock_provider,
    )
    rows = manifest.get("sentences", [])
    blocked_externalization_violations = 0
    raw_sentence_leak_violations = 0
    normalized_queries = [re.sub(r"\s+", " ", str(q).strip().lower()) for q in mock_provider.queries]

    for row in rows if isinstance(rows, list) else []:
        if not isinstance(row, dict):
            continue
        privacy = row.get("privacy", {})
        safe = bool(privacy.get("safe_for_external_search")) if isinstance(privacy, dict) else False
        used_external = bool(row.get("used_external_search"))
        if not safe and used_external:
            blocked_externalization_violations += 1
        sentence_norm = re.sub(r"\s+", " ", str(row.get("sentence_text", "")).strip().lower())
        if sentence_norm and sentence_norm in normalized_queries:
            raw_sentence_leak_violations += 1

    return {
        "probe_sentences": len(rows) if isinstance(rows, list) else 0,
        "queries_emitted": len(mock_provider.queries),
        "blocked_externalization_violations": blocked_externalization_violations,
        "raw_sentence_leak_violations": raw_sentence_leak_violations,
        "passed": blocked_externalization_violations == 0 and raw_sentence_leak_violations == 0,
    }


def _is_fixture_candidate_paragraph(text: str, section: str) -> bool:
    raw = str(text or "").strip()
    if not raw:
        return False
    if section in {"front_matter", "references", "header_footer"}:
        return False
    low = raw.lower()
    if any(mark in low for mark in ["start of picture text", "end of picture text", "<br>"]):
        return False
    if raw.lstrip().startswith(("-", "•", "* ")):
        return False
    if low.startswith(("fig.", "figure ", "table ", "scheme ", "supplementary")):
        return False
    if low.startswith(("**fig.", "**figure ", "**table ", "**scheme ")):
        return False
    if re.match(r"^\(?\d{1,3}\)\s+", raw):
        return False
    if re.match(r"^\d{1,3}\.\s+", raw):
        return False
    if re.search(r"\b(et al\.|nat\.|science|j\. org\.|tetrahedron|drug discovery|acc\. chem\. res\.)\b", low):
        return False
    if any(k in low for k in ["springer", "wiley", "doi", "pmid"]):
        return False
    if re.search(r"\(\d{4}[a-z]?\)\.?$", raw) and re.search(r"\b\d{1,4}\s*[,:\u2013-]\s*\d{1,5}\b", raw):
        return False
    if re.search(r"\b\d{4}\b", raw) and raw.count(";") >= 2:
        return False
    if raw.count(";") >= 3:
        return False
    if re.search(r"\b\d{1,4}\s*[,:\u2013-]\s*\d{1,5}\.?$", raw):
        return False
    words = raw.split()
    if len(words) < 12 or len(words) > 70:
        return False
    return True


def _build_existing_comment_fixture(source_docx: Path, fixture_docx: Path) -> dict[str, Any]:
    fixture_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_docx, fixture_docx)
    doc = Document(str(fixture_docx))
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    section_by_idx = _build_section_index_map(paragraphs)
    chosen_indices: list[int] = []
    preferred = ["introduction", "methods", "results", "discussion", "conclusions", "abstract", "body"]
    for sec in preferred:
        for idx, text in enumerate(paragraphs):
            if idx in chosen_indices:
                continue
            if section_by_idx.get(idx, "body") != sec:
                continue
            if _is_fixture_candidate_paragraph(text, sec):
                chosen_indices.append(idx)
                if len(chosen_indices) >= 3:
                    break
        if len(chosen_indices) >= 3:
            break

    comment_templates = [
        "Please split this sentence and tighten the flow without changing meaning.",
        "This claim sounds strong. Qualify scope and add direct support if available.",
        "Wording is clunky here; suggest a concise local rewrite.",
    ]
    inserted = 0
    for pos, pidx in enumerate(chosen_indices):
        if pidx < 0 or pidx >= len(doc.paragraphs):
            continue
        paragraph = doc.paragraphs[pidx]
        if not paragraph.runs:
            paragraph.add_run(paragraph.text if paragraph.text else " ")
        comment = doc.comments.add_comment(
            text=comment_templates[pos % len(comment_templates)],
            author="Eval Reviewer",
            initials="ER",
        )
        paragraph.runs[0].mark_comment_range(paragraph.runs[-1], comment.comment_id)
        inserted += 1
    doc.save(str(fixture_docx))
    return {
        "fixture_docx": str(fixture_docx),
        "inserted_comments": inserted,
        "target_paragraph_indices": chosen_indices,
    }


def _responder_stats(manifest: dict[str, Any]) -> dict[str, int]:
    summary = manifest.get("summary", {}) if isinstance(manifest, dict) else {}
    rows = manifest.get("responses", []) if isinstance(manifest, dict) else []
    concrete = 0
    vague = 0
    for row in rows if isinstance(rows, list) else []:
        if not isinstance(row, dict):
            continue
        rtype = str(row.get("response_type", "")).strip().lower()
        proposed = str(row.get("proposed_response", "")).strip()
        if rtype in {"text_fix", "response_strategy", "already_addressed"} and proposed:
            concrete += 1
        if rtype == "needs_clarification":
            vague += 1
    return {
        "source_comments_read_count": int(summary.get("existing_comments_detected", 0) or 0),
        "responses_generated_count": int(summary.get("responses_generated", 0) or 0),
        "concrete_response_count": concrete,
        "needs_clarification_count": vague,
    }


def _run_case(
    case: dict[str, str],
    out_root: Path,
    provider: OllamaProvider | None,
    judge_model: str | None,
    run_id: str,
) -> tuple[dict[str, Any], dict[str, Any]]:
    source_docx = Path(case["source_docx"])
    review_json = Path(case["review_json"])
    if not source_docx.exists():
        raise FileNotFoundError(f"Missing source docx: {source_docx}")
    if not review_json.exists():
        raise FileNotFoundError(f"Missing review json: {review_json}")
    out_dir = out_root / case["project_id"]
    out_dir.mkdir(parents=True, exist_ok=True)

    review_payload = json.loads(review_json.read_text(encoding="utf-8"))
    review_obj = ReviewObj(review_payload)
    annotation = build_annotated_manuscript_output(
        source_path=source_docx,
        doc=_build_doc(source_docx),
        review=review_obj,
        output_dir=out_dir,
        project_id=case["project_id"],
        run_id=run_id,
        provider=provider,
        model=None,
        rewrite_model=None,
        comment_audit_model=judge_model,
        suggestion_audit_model=judge_model,
        reflection_model=None,
        adjudication_model=None,
        supporting_cards=None,
        claim_check_external_enabled=False,
        timeout_seconds=180,
    )
    manifest_path = out_dir / "manuscript_comment_manifest.json"
    manifest_path.write_text(json.dumps(annotation, indent=2), encoding="utf-8")

    metrics = compute_comment_quality_metrics(annotation)
    metrics["privacy_search_safety"]["mock_probe"] = _run_privacy_probe(source_docx)
    fixture_meta: dict[str, Any] | None = None
    fixture_responder_manifest: dict[str, Any] | None = None
    base_responder_summary = _responder_stats(annotation.get("comment_response_manifest", {}))
    if base_responder_summary["source_comments_read_count"] == 0:
        fixture_docx = out_dir / "fixtures" / f"{source_docx.stem}_existing_comments_fixture.docx"
        fixture_meta = _build_existing_comment_fixture(source_docx, fixture_docx)
        if int(fixture_meta.get("inserted_comments", 0) or 0) > 0:
            fixture_paragraphs = _analysis_paragraphs(fixture_docx)
            fixture_sections = _build_section_index_map(fixture_paragraphs)
            fixture_responder_manifest = run_existing_comment_responder(
                source_docx=fixture_docx,
                paragraphs=fixture_paragraphs,
                section_by_idx=fixture_sections,
                provider=provider,
                model=judge_model,
                timeout_seconds=120,
            )
            (out_dir / "comment_response_fixture_manifest.json").write_text(
                json.dumps(fixture_responder_manifest, indent=2), encoding="utf-8"
            )
            (out_dir / "comment_response_fixture_manifest.md").write_text(
                render_comment_response_manifest_markdown(fixture_responder_manifest), encoding="utf-8"
            )
            (out_dir / "comment_response_fixture_metadata.json").write_text(
                json.dumps(fixture_meta, indent=2), encoding="utf-8"
            )
    if fixture_responder_manifest is not None:
        fixture_stats = _responder_stats(fixture_responder_manifest)
        coverage = metrics.get("existing_comment_responder_coverage", {})
        if isinstance(coverage, dict):
            coverage["fixture_source_comments_read_count"] = fixture_stats["source_comments_read_count"]
            coverage["fixture_responses_generated_count"] = fixture_stats["responses_generated_count"]
            coverage["fixture_concrete_response_count"] = fixture_stats["concrete_response_count"]
            coverage["fixture_needs_clarification_count"] = fixture_stats["needs_clarification_count"]
            coverage["effective_source_comments_read_count"] = max(
                int(coverage.get("source_comments_read_count", 0) or 0),
                fixture_stats["source_comments_read_count"],
            )
            coverage["effective_concrete_response_count"] = max(
                int(coverage.get("concrete_response_count", 0) or 0),
                fixture_stats["concrete_response_count"],
            )
            coverage["effective_needs_clarification_count"] = max(
                int(coverage.get("needs_clarification_count", 0) or 0),
                fixture_stats["needs_clarification_count"],
            )
            coverage["effective_concrete_response_rate"] = round(
                coverage["effective_concrete_response_count"] / max(1, coverage["effective_source_comments_read_count"]),
                4,
            )
    else:
        coverage = metrics.get("existing_comment_responder_coverage", {})
        if isinstance(coverage, dict):
            coverage["effective_source_comments_read_count"] = int(coverage.get("source_comments_read_count", 0) or 0)
            coverage["effective_concrete_response_count"] = int(coverage.get("concrete_response_count", 0) or 0)
            coverage["effective_needs_clarification_count"] = int(coverage.get("needs_clarification_count", 0) or 0)
            coverage["effective_concrete_response_rate"] = round(
                coverage["effective_concrete_response_count"] / max(1, coverage["effective_source_comments_read_count"]),
                4,
            )
    (out_dir / "comment_quality_metrics.json").write_text(json.dumps(metrics, indent=2), encoding="utf-8")
    claim_manifest = annotation.get("sentence_claim_check", {}) if isinstance(annotation, dict) else {}
    claim_projection = claim_manifest.get("comment_projection", {}) if isinstance(claim_manifest, dict) else {}
    surfaced = [
        c
        for c in annotation.get("comment_targets", [])
        if isinstance(c, dict) and bool(c.get("from_claim_check"))
    ] if isinstance(annotation, dict) else []
    suppressed_examples = claim_projection.get("suppressed_examples", {}) if isinstance(claim_projection, dict) else {}
    lines = [
        "# Claim-Check Examples",
        "",
        "## Surfaced Claim-Check Comments",
    ]
    if surfaced:
        for row in surfaced[:8]:
            lines.extend(
                [
                    f"- {row.get('comment_id', 'unknown')} | verdict={row.get('claim_check_verdict', 'unknown')} | paragraph={row.get('paragraph_index', 'n/a')}",
                    f"  - anchor: {str(row.get('anchor_text', '')).strip()[:220]}",
                    f"  - critique: {str(row.get('critique', '')).strip()[:180]}",
                    f"  - action: {str(row.get('suggested_revision', '')).strip()[:180]}",
                ]
            )
    else:
        lines.append("- none")
    lines.extend(["", "## Suppressed Sentence IDs By Reason"])
    if isinstance(suppressed_examples, dict) and suppressed_examples:
        for reason, values in sorted(suppressed_examples.items(), key=lambda kv: str(kv[0])):
            if not isinstance(values, list):
                continue
            lines.append(f"- {reason}: {', '.join(str(v) for v in values[:8])}")
    else:
        lines.append("- none")
    (out_dir / "claim_check_examples.md").write_text("\n".join(lines) + "\n", encoding="utf-8")
    claim_model = ((annotation.get("sentence_claim_check") or {}).get("model") or {}) if isinstance(annotation, dict) else {}
    style_model = ((annotation.get("style_clarity_pass") or {}).get("model") or {}) if isinstance(annotation, dict) else {}
    responder_model = ((annotation.get("comment_response_manifest") or {}).get("model") or {}) if isinstance(annotation, dict) else {}
    case_meta = {
        "project_id": case["project_id"],
        "claim_model": claim_model,
        "style_model": style_model,
        "responder_model": responder_model,
        "degraded_mode": bool(
            claim_model.get("fallback_used")
            or style_model.get("fallback_used")
            or responder_model.get("fallback_used")
        ),
        "fixture_comment_injection": fixture_meta or {"inserted_comments": 0},
        "fixture_responder_used": bool(fixture_responder_manifest is not None),
    }
    (out_dir / "claim_check_case_run_metadata.json").write_text(json.dumps(case_meta, indent=2), encoding="utf-8")
    return metrics, case_meta


def _render_summary_markdown(metrics_by_case: dict[str, dict[str, Any]]) -> str:
    lines = [
        "# Comment Quality Evaluation Summary",
        "",
        "| Project | Avg anchor | Paragraph-wide local rate | Sentence-local overlong | Clause overbroad | No-op rewrites | Trivial rewrites | Materially better rewrites | Avg comment len | Near-dup pair rate | Overlap rate | Overclaim rate | Unclear-support rate | Surfaced claim comments | Style surfaced | Safe externalizable | Blocked externalizable |",
        "|---|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|---:|",
    ]
    for project_id, metrics in metrics_by_case.items():
        anchor = metrics.get("anchor_localization", {})
        rewrite = metrics.get("rewrite_usefulness", {})
        concision = metrics.get("comment_concision", {})
        dedupe = metrics.get("deduplication_quality", {})
        claim = metrics.get("claim_check_coverage", {})
        privacy = metrics.get("privacy_search_safety", {})
        style = metrics.get("style_pass_outcomes", {})
        lines.append(
            "| "
            + f"{project_id} | "
            + f"{anchor.get('avg_anchor_length_words', 0)} | "
            + f"{anchor.get('paragraph_wide_local_anchor_rate', 0)} | "
            + f"{anchor.get('sentence_local_exceed_one_sentence_rate', 0)} | "
            + f"{anchor.get('clause_local_too_broad_rate', 0)} | "
            + f"{rewrite.get('no_op_rewrite_rate', 0)} | "
            + f"{rewrite.get('trivial_rewrite_rate', 0)} | "
            + f"{rewrite.get('materially_better_surviving_rate', 0)} | "
            + f"{concision.get('avg_comment_length_words', 0)} | "
            + f"{dedupe.get('near_duplicate_pair_rate', 0)} | "
            + f"{dedupe.get('overlapping_comment_rate', 0)} | "
            + f"{claim.get('overclaim_rate', 0)} | "
            + f"{claim.get('unclear_support_rate', 0)} | "
            + f"{claim.get('surfaced_comment_count', 0)} | "
            + f"{style.get('style_comments_surfaced_count', 0)} | "
            + f"{privacy.get('safe_externalization_count', 0)} | "
            + f"{privacy.get('blocked_externalization_count', 0)} |"
        )
    return "\n".join(lines) + "\n"


def _render_comparison_markdown(compare_by_case: dict[str, dict[str, dict[str, float | int | None]]]) -> str:
    lines = [
        "# Comment Quality Comparison",
        "",
    ]
    for project_id, rows in compare_by_case.items():
        lines.extend(
            [
                f"## {project_id}",
                "",
                "| Metric | Before | After | Delta |",
                "|---|---:|---:|---:|",
            ]
        )
        for key in COMPARISON_KEYS:
            row = rows.get(key, {})
            lines.append(
                f"| {key} | {row.get('before')} | {row.get('after')} | {row.get('delta')} |"
            )
        lines.append("")
    return "\n".join(lines)


def main() -> None:
    parser = argparse.ArgumentParser(description="Run comment quality evaluation on allowed projects.")
    parser.add_argument("--out-root", required=True, help="Output directory for evaluation artifacts.")
    parser.add_argument("--label", default="", help="Optional label for metadata.")
    parser.add_argument("--baseline-dir", default="", help="Optional baseline eval directory for comparison.")
    parser.add_argument("--use-provider", action="store_true", help="Enable local provider if available.")
    parser.add_argument("--judge-model", default="", help="Judge model; gemma4:31b preferred when available.")
    args = parser.parse_args()

    out_root = Path(args.out_root)
    out_root.mkdir(parents=True, exist_ok=True)
    run_id = f"comment_quality_eval_{out_root.name}"

    provider: OllamaProvider | None = None
    judge_model = args.judge_model.strip() or None
    degraded_mode = True
    if args.use_provider:
        candidate = OllamaProvider()
        ok, _ = candidate.health()
        if ok:
            available = set(candidate.list_models())
            if judge_model and judge_model not in available:
                provider = None
                judge_model = None
            else:
                provider = candidate
                degraded_mode = False

    metrics_by_case: dict[str, dict[str, Any]] = {}
    case_meta_rows: list[dict[str, Any]] = []
    for case in ALLOWED_CASES:
        metrics, case_meta = _run_case(
            case=case,
            out_root=out_root,
            provider=provider,
            judge_model=judge_model,
            run_id=run_id,
        )
        metrics_by_case[case["project_id"]] = metrics
        case_meta_rows.append(case_meta)

    summary_json_path = out_root / "metrics_summary.json"
    summary_json_path.write_text(json.dumps(metrics_by_case, indent=2), encoding="utf-8")
    summary_md_path = out_root / "metrics_summary.md"
    summary_md_path.write_text(_render_summary_markdown(metrics_by_case), encoding="utf-8")

    compare_payload: dict[str, Any] = {}
    baseline_dir = Path(args.baseline_dir).expanduser() if args.baseline_dir else None
    if baseline_dir:
        baseline_metrics_path = baseline_dir / "metrics_summary.json"
        if baseline_metrics_path.exists():
            baseline_metrics = json.loads(baseline_metrics_path.read_text(encoding="utf-8"))
            compare_by_case: dict[str, dict[str, dict[str, float | int | None]]] = {}
            for case in ALLOWED_CASES:
                pid = case["project_id"]
                before = baseline_metrics.get(pid, {})
                after = metrics_by_case.get(pid, {})
                compare_by_case[pid] = compare_metric_dicts(before, after, COMPARISON_KEYS)
            compare_payload = compare_by_case
            (out_root / "metrics_comparison_vs_baseline.json").write_text(json.dumps(compare_by_case, indent=2), encoding="utf-8")
            (out_root / "metrics_comparison_vs_baseline.md").write_text(_render_comparison_markdown(compare_by_case), encoding="utf-8")

    metadata = {
        "generated_at": datetime.now(timezone.utc).isoformat(),
        "out_root": str(out_root),
        "label": args.label,
        "allowed_projects": [case["project_id"] for case in ALLOWED_CASES],
        "use_provider": bool(provider is not None),
        "judge_model": judge_model,
        "degraded_mode": degraded_mode or any(bool(row.get("degraded_mode")) for row in case_meta_rows),
        "baseline_dir": str(baseline_dir) if baseline_dir else None,
        "comparison_generated": bool(compare_payload),
        "case_model_metadata": case_meta_rows,
    }
    (out_root / "eval_run_metadata.json").write_text(json.dumps(metadata, indent=2), encoding="utf-8")
    print(str(out_root))


if __name__ == "__main__":
    main()
