import os
from pathlib import Path
from docx import Document
from ai_reviewer.tools.docx_tools import create_commented_docx_copy, inspect_docx_annotation_state

def create_synthetic_docs():
    out_dir = Path("audits/docx_native_fixtures")
    out_dir.mkdir(parents=True, exist_ok=True)
    
    # 1. Clean Native DOCX
    clean_path = out_dir / "clean_native.docx"
    doc = Document()
    doc.add_heading("Synthetic Paper", level=1)
    doc.add_heading("Abstract", level=2)
    doc.add_paragraph("This is a clean abstract with a strong claim. It has no comments.")
    doc.add_heading("Introduction", level=2)
    doc.add_paragraph("This is an introduction paragraph. It sets up the problem.")
    doc.save(str(clean_path))

    # 2. Native DOCX with human comments
    # Python-docx doesn't natively support adding comments easily without our own tool, so we will simulate it by adding an "AI-Reviewer" comment and changing the author if possible, or just using AI-Reviewer comment as a proxy for 'existing comments'.
    commented_path = out_dir / "human_commented.docx"
    create_commented_docx_copy(
        source_docx=clean_path,
        output_docx=commented_path,
        comments=[
            {
                "paragraph_index": 2,
                "issue_type": "human review",
                "severity": "medium",
                "critique": "I think this claim is a bit strong.",
                "suggested_revision": "Soften the claim.",
                "rationale": "Just a thought.",
                "anchor_text": "This is a clean abstract with a strong claim."
            }
        ],
        author="Reviewer 1",
        initials="R1"
    )

    # 3. Native DOCX with AI-Reviewer comments
    ai_commented_path = out_dir / "ai_commented.docx"
    create_commented_docx_copy(
        source_docx=clean_path,
        output_docx=ai_commented_path,
        comments=[
            {
                "paragraph_index": 4,
                "issue_type": "clarity",
                "severity": "low",
                "critique": "This sentence is too short.",
                "suggested_revision": "Add more detail.",
                "rationale": "Heuristic.",
                "anchor_text": "It sets up the problem."
            }
        ],
        author="AI-Reviewer",
        initials="AIR",
        comment_tag="Run 1"
    )

    print("Created synthetic test docs in audits/docx_native_fixtures/")

if __name__ == "__main__":
    create_synthetic_docs()
