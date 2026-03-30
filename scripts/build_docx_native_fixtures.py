from __future__ import annotations

import json
import shutil
from pathlib import Path

from docx import Document


REPO_ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = REPO_ROOT / "audits" / "docx_native_fixtures" / "generated"

PROJECTS = {
    "project1": {
        "project_id": "20260325163524_test-existingphactorpaper",
        "native_name": "project1_clean_native.docx",
        "comment_targets": {
            "light": [
                (7, "Editor A", "EA", "Abstract claim feels broad; tie the success claim to tested conditions."),
                (13, "Lab Reviewer", "LR", "Methods list item is understandable, but the operational constraint should be more explicit."),
                (94, "Journal Editor", "JE", "Discussion sentence sounds broader than the evidence shown here."),
            ],
            "heavy": [
                (7, "Editor A", "EA", "Abstract claim feels broad; tie the success claim to tested conditions."),
                (10, "Lab Reviewer", "LR", "This transition sentence is overloaded and should be split."),
                (13, "Methods Reviewer", "MR", "Clarify the experimental constraint rather than implying it."),
                (14, "Methods Reviewer", "MR", "The software handoff step needs a cleaner statement."),
                (42, "Results Reviewer", "RR", "Avoid overstating the better performance claim."),
                (94, "Journal Editor", "JE", "This discussion sentence overgeneralizes from the tested examples."),
            ],
        },
    },
    "project2": {
        "project_id": "20260327051312_miniaturization_d2b",
        "native_name": "project2_clean_native.docx",
        "comment_targets": {
            "light": [
                (6, "Editor A", "EA", "Introduction framing is too dense for the opening claim."),
                (36, "Methods Reviewer", "MR", "Methods sentence needs the comparison point earlier."),
                (80, "Results Reviewer", "RR", "This paragraph needs tighter claim calibration."),
            ],
            "heavy": [
                (6, "Editor A", "EA", "Introduction framing is too dense for the opening claim."),
                (7, "Editor A", "EA", "This bridge sentence is still too broad."),
                (18, "Results Reviewer", "RR", "This result sentence carries too many constraints."),
                (19, "Results Reviewer", "RR", "The catalyst claim needs scope language."),
                (36, "Methods Reviewer", "MR", "Methods sentence needs the comparison point earlier."),
                (67, "Discussion Editor", "DE", "This transition to Boc deprotection needs a clearer boundary."),
                (80, "Results Reviewer", "RR", "This paragraph needs tighter claim calibration."),
            ],
        },
    },
}


def _add_comment(doc: Document, paragraph_index: int, author: str, initials: str, text: str) -> None:
    if paragraph_index >= len(doc.paragraphs):
        return
    paragraph = doc.paragraphs[paragraph_index]
    if not paragraph.runs:
        paragraph.add_run(" ")
    first_run = paragraph.runs[0]
    last_run = paragraph.runs[-1]
    comment = doc.comments.add_comment(text=text, author=author, initials=initials)
    first_run.mark_comment_range(last_run, comment.comment_id)


def _build_commented_fixture(source: Path, target: Path, targets: list[tuple[int, str, str, str]]) -> None:
    doc = Document(str(source))
    for paragraph_index, author, initials, text in targets:
        _add_comment(doc, paragraph_index, author, initials, text)
    target.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(target))


def _latest_path(paths: list[Path]) -> Path:
    if not paths:
        raise FileNotFoundError("No candidate paths found for fixture generation.")
    return sorted(paths, key=lambda p: str(p))[-1]


def _resolve_native_source(project_id: str, native_name: str) -> Path:
    candidates = list((REPO_ROOT / "projects" / project_id / "materials" / "managed").rglob(native_name))
    if candidates:
        return _latest_path(candidates)
    raise FileNotFoundError(f"No native DOCX source found for {project_id}: {native_name}")


def _resolve_prior_comment_source(project_id: str) -> Path:
    candidates = list((REPO_ROOT / "projects" / project_id / "runs").rglob("reviewed_manuscript_with_comments.docx"))
    if candidates:
        return _latest_path(candidates)
    candidates = list((REPO_ROOT / "projects" / project_id / "runs").rglob("surrogate_manuscript_from_pdf_with_comments.docx"))
    return _latest_path(candidates)


def _resolve_prior_suggested_source(project_id: str) -> Path:
    candidates = list((REPO_ROOT / "projects" / project_id / "runs").rglob("reviewed_manuscript_with_suggested_changes.docx"))
    if candidates:
        return _latest_path(candidates)
    candidates = list((REPO_ROOT / "projects" / project_id / "runs").rglob("surrogate_manuscript_from_pdf_with_suggested_changes.docx"))
    return _latest_path(candidates)


def build_fixtures() -> dict:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    manifest: dict[str, dict] = {"output_dir": str(OUT_DIR), "fixtures": {}}
    for key, cfg in PROJECTS.items():
        project_id = cfg["project_id"]
        clean_source = _resolve_native_source(project_id, cfg["native_name"])
        prior_commented_source = _resolve_prior_comment_source(project_id)
        prior_suggested_source = _resolve_prior_suggested_source(project_id)

        project_out = OUT_DIR / project_id
        project_out.mkdir(parents=True, exist_ok=True)

        clean = project_out / f"{key}_clean_native.docx"
        light = project_out / f"{key}_commented_light.docx"
        heavy = project_out / f"{key}_commented_heavy.docx"
        prior_comments = project_out / f"{key}_prior_ai_commented.docx"
        prior_suggested = project_out / f"{key}_prior_ai_suggested.docx"

        shutil.copy2(clean_source, clean)
        _build_commented_fixture(clean_source, light, cfg["comment_targets"]["light"])
        _build_commented_fixture(clean_source, heavy, cfg["comment_targets"]["heavy"])
        shutil.copy2(prior_commented_source, prior_comments)
        shutil.copy2(prior_suggested_source, prior_suggested)

        manifest["fixtures"][project_id] = {
            "clean_native": str(clean),
            "commented_light": str(light),
            "commented_heavy": str(heavy),
            "prior_ai_commented": str(prior_comments),
            "prior_ai_suggested": str(prior_suggested),
            "sources": {
                "clean_source": str(clean_source),
                "prior_commented_source": str(prior_commented_source),
                "prior_suggested_source": str(prior_suggested_source),
            },
        }
    manifest_path = OUT_DIR / "fixture_manifest.json"
    manifest_path.write_text(json.dumps(manifest, indent=2), encoding="utf-8")
    return manifest


if __name__ == "__main__":
    built = build_fixtures()
    print(json.dumps(built, indent=2))
