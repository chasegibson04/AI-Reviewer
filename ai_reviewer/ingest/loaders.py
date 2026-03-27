from __future__ import annotations

from datetime import datetime, timezone
import re
from pathlib import Path

from docx import Document as DocxDocument
from pypdf import PdfReader
import pymupdf

try:
    import pymupdf4llm
except Exception:  # pragma: no cover
    pymupdf4llm = None

from ai_reviewer.paths import REPO_ROOT
from ai_reviewer.ingest.types import ParseFailure, ParsedDocument
from ai_reviewer.tools.pdf_tools import parse_pdf_structured
from ai_reviewer.tools.docx_tools import parse_docx_structured

SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".markdown", ".docx", ".tex"}


def _warn_low_quality(text: str, warnings: list[str]) -> None:
    stripped = text.strip()
    if len(stripped) < 1000:
        warnings.append("Tiny extraction length; parser may have failed.")

    bad_ratio = 0.0
    if text:
        bad_ratio = sum(1 for c in text if c in {"\uFFFD", "?", "\u00ad"}) / max(len(text), 1)
    if bad_ratio > 0.08:
        warnings.append("Text contains repeated broken OCR or encoding artifacts.")

    if re.search(r"[A-Za-z]\s+[A-Za-z]\s+[A-Za-z]\s+[A-Za-z]\s+[A-Za-z]", text):
        warnings.append("Detected character-spaced OCR artifacts.")

    if "ﬁ" in text or "ﬂ" in text:
        warnings.append("Detected ligatures; downstream tokenization may be affected.")

    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    if lines:
        repeats = len(lines) - len(set(lines))
        if repeats > max(8, int(0.2 * len(lines))):
            warnings.append("Large repeated-text ratio suggests duplicated pages/extraction loops.")

    bib_hits = len(re.findall(r"\[[0-9]{1,3}\]", text)) + len(re.findall(r"\b(doi|arxiv|proceedings)\b", text.lower()))
    if bib_hits > 60:
        warnings.append("Extraction appears bibliography-heavy; core sections may be underrepresented.")


def guess_headings(text: str) -> list[str]:
    headings: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        if re.match(r"^(\d+(\.\d+)*)\s+[A-Z]", stripped):
            headings.append(stripped)
        elif stripped.startswith("#"):
            headings.append(stripped.lstrip("# "))
        elif len(stripped) < 80 and stripped.isupper() and len(stripped.split()) <= 14:
            headings.append(stripped.title())
    return headings[:300]


def _read_pdf(path: Path, warnings: list[str]) -> tuple[str, int | None, str]:
    try:
        parsed = parse_pdf_structured(path)
        for w in parsed.get("warnings", []):
            warnings.append(str(w))
        warnings.append(
            f"pdf_structure:figures={parsed.get('figure_mentions', 0)},tables={parsed.get('table_mentions', 0)},citations={len(parsed.get('citation_markers', []))}"
        )
        text = parsed.get("text", "")
        page_count = parsed.get("page_count")
        engine = parsed.get("tool", "pdf-tools")
        if len(text.strip()) < 200:
            warnings.append("Likely scanned PDF or extraction failure.")
        return text, page_count, engine
    except Exception as exc:
        warnings.append(f"structured_pdf_parse_failed:{exc}")

    text = ""
    engine = "none"
    if pymupdf4llm is not None:
        try:
            text = pymupdf4llm.to_markdown(str(path))
            engine = "pymupdf4llm"
        except Exception as exc:
            warnings.append(f"pymupdf4llm parse failure: {exc}")
    if len(text.strip()) < 200:
        try:
            doc = pymupdf.open(str(path))
            text = "\n".join(page.get_text() for page in doc)
            engine = "pymupdf"
        except Exception as exc:
            warnings.append(f"pymupdf parse failure: {exc}")
    page_count = None
    try:
        reader = PdfReader(str(path))
        page_count = len(reader.pages)
        if len(text.strip()) < 200:
            text = "\n".join((p.extract_text() or "") for p in reader.pages)
            engine = "pypdf"
    except Exception as exc:
        warnings.append(f"pypdf parse failure: {exc}")
    return text, page_count, engine


def _read_docx(path: Path) -> tuple[str, str]:
    try:
        parsed = parse_docx_structured(path)
        return parsed.get("text", ""), "python-docx+structured"
    except Exception:
        doc = DocxDocument(str(path))
        return "\n".join(p.text for p in doc.paragraphs), "python-docx"


def clean_text(text: str) -> str:
    text = text.replace("\x00", "")
    text = text.replace("\u00ad", "")
    text = text.replace("\ufb01", "fi").replace("\ufb02", "fl")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()


def _relative_path(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(REPO_ROOT.resolve()))
    except Exception:
        return str(path)


def parse_file(path: Path) -> ParsedDocument:
    warnings: list[str] = []
    suffix = path.suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise ValueError(f"Unsupported file type: {path.suffix}")

    raw_text = ""
    page_count: int | None = None
    parse_engine = "plain-text"

    if suffix == ".pdf":
        raw_text, page_count, parse_engine = _read_pdf(path, warnings)
    elif suffix in {".txt", ".md", ".markdown", ".tex"}:
        raw_text = path.read_text(encoding="utf-8", errors="replace")
        parse_engine = "utf8-text"
    elif suffix == ".docx":
        raw_text, parse_engine = _read_docx(path)

    cleaned = clean_text(raw_text)
    headings = guess_headings(cleaned)
    if not headings:
        warnings.append("No section headings detected.")

    _warn_low_quality(cleaned, warnings)

    if not cleaned:
        warnings.append("Empty extraction result.")

    return ParsedDocument(
        source_path=path,
        source_path_abs=path.resolve(),
        source_path_rel=_relative_path(path),
        file_size_bytes=path.stat().st_size,
        document_type=suffix.lstrip("."),
        parse_engine=parse_engine,
        ingest_timestamp=datetime.now(timezone.utc).isoformat(),
        raw_text=raw_text,
        cleaned_text=cleaned,
        headings=headings,
        page_count=page_count,
        parse_warnings=warnings,
    )


def collect_paths(path: Path) -> list[Path]:
    if path.is_file():
        return [path]
    if not path.exists():
        raise FileNotFoundError(f"Input path not found: {path}")

    discovered: list[Path] = []
    for ext in SUPPORTED_EXTENSIONS:
        discovered.extend(path.rglob(f"*{ext}"))
    return sorted(set(discovered))


def parse_path_with_failures(path: Path, continue_on_error: bool = False) -> tuple[list[ParsedDocument], list[ParseFailure]]:
    docs: list[ParsedDocument] = []
    failures: list[ParseFailure] = []
    for src in collect_paths(path):
        try:
            docs.append(parse_file(src))
        except Exception as exc:
            failures.append(ParseFailure(source_path=src, error=str(exc)))
            if not continue_on_error:
                break
    return docs, failures


def parse_path(path: Path) -> list[ParsedDocument]:
    docs, failures = parse_path_with_failures(path, continue_on_error=False)
    if failures:
        raise RuntimeError("Some files failed to parse:\n" + "\n".join(f"{f.source_path}: {f.error}" for f in failures))
    return docs
