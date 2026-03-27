from __future__ import annotations

from pathlib import Path

from docx import Document


def write_markdown_as_docx(markdown_text: str, output_path: Path) -> None:
    doc = Document()
    lines = markdown_text.splitlines()
    for raw in lines:
        line = raw.rstrip()
        if not line:
            doc.add_paragraph("")
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        doc.add_paragraph(line)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
