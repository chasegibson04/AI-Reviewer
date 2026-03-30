from __future__ import annotations

from pathlib import Path
from typing import Any
import json
import re
from uuid import uuid4

from docx import Document

from ai_reviewer.ingest.types import ParsedDocument
from ai_reviewer.models.base import Provider, ChatRequest
from ai_reviewer.review.repair import extract_json_candidate
from ai_reviewer.tools.docx_tools import (
    create_commented_docx_copy,
    create_docx_from_plain_text,
    create_suggested_changes_docx,
    get_docx_paragraph_texts,
    inspect_docx_annotation_state,
    normalize_review_artifact_text,
    validate_commented_docx,
    validate_suggested_changes_docx,
)


def detect_source_mode(path: Path) -> dict[str, Any]:
    suffix = path.suffix.lower()
    if suffix == ".docx":
        state = inspect_docx_annotation_state(path) if path.exists() else {"annotation_state": "unknown_docx_state"}
        return {
            "mode": "original_docx",
            "base_type": "docx",
            "annotation_state": state.get("annotation_state"),
            "docx_annotation_state": state,
        }
    if suffix == ".pdf":
        return {"mode": "pdf_only_surrogate", "base_type": "pdf"}
    return {"mode": "surrogate_other_source", "base_type": suffix.lstrip(".")}


def _analysis_paragraphs(base_docx: Path) -> list[str]:
    return [text.strip() for text in get_docx_paragraph_texts(base_docx, normalize_review_artifacts=True)]


def _split_sentences(text: str) -> list[str]:
    return [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if s.strip()]


def _normalize_section_target(section: str | None) -> str | None:
    low = re.sub(r"[^a-zA-Z ]", " ", str(section or "")).lower()
    low = re.sub(r"\s+", " ", low).strip()
    if not low:
        return None
    if "abstract" in low:
        return "abstract"
    if "intro" in low or "background" in low:
        return "introduction"
    if any(k in low for k in ["method", "experimental", "procedure"]):
        return "methods"
    if "result" in low or "finding" in low:
        return "results"
    if "discussion" in low:
        return "discussion"
    if "conclusion" in low:
        return "conclusions"
    return None


def _calibrate_claim_sentence(sentence: str) -> str:
    revised = sentence
    replacements = [
        (r"\bevery instance tried\b", "at least one successful condition in each case study"),
        (r"\bon the first attempt\b", "in the initial screen"),
        (r"\bfirst attempt\b", "initial screen"),
        (r"\bmodest to excellent\b", "modest to high"),
        (r"\bclearly demonstrates\b", "suggests"),
        (r"\bdemonstrates\b", "suggests"),
        (r"\bproves\b", "supports"),
        (r"\ball\b", "the tested"),
        (r"\balways\b", "often"),
        (r"\bnever\b", "did not"),
    ]
    for pattern, replacement in replacements:
        revised = re.sub(pattern, replacement, revised, flags=re.IGNORECASE)
    return revised


def _build_sentence_level_candidates(base_docx: Path, max_comments: int) -> list[dict[str, Any]]:
    paragraphs = _analysis_paragraphs(base_docx)
    section_by_idx = _build_section_index_map(paragraphs)
    candidates: list[dict[str, Any]] = []
    per_section: dict[str, int] = {}
    for pidx, text in enumerate(paragraphs):
        if len(candidates) >= max_comments:
            break
        if not text or len(text.split()) < 16:
            continue
        low = text.lower()
        section = section_by_idx.get(pidx, "body")
        if section in {"front_matter", "references", "header_footer"}:
            continue
        if _is_placeholder_paragraph(low) or _is_heading_paragraph(text) or _is_front_or_back_matter_text(low):
            continue
        if section == "abstract" and per_section.get(section, 0) >= 1:
            continue
        sentences = _split_sentences(text)
        for sentence in sentences:
            if len(candidates) >= max_comments:
                break
            words = sentence.split()
            if len(words) < 12:
                continue
            slow = sentence.lower()
            entry: dict[str, Any] | None = None
            if section in {"abstract", "results", "discussion", "conclusions", "introduction"} and any(
                phrase in slow
                for phrase in ["first attempt", "every instance", "all ", "always", "never", "demonstrates", "clearly demonstrates", "proves"]
            ):
                entry = {
                    "paragraph_index": pidx,
                    "issue_type": "evidence/overclaim concern",
                    "severity": "high",
                    "critique": f'This sentence reads broader than the evidence shown here: "{sentence[:220]}". Narrow the scope or name the tested condition explicitly.',
                    "suggested_revision": f"Proposed edit: {_calibrate_claim_sentence(sentence)}",
                    "rationale": "Sentence-level claim calibration grounded in manuscript wording.",
                    "locked_paragraph": True,
                    "anchor_text": sentence,
                    "span_sentence": sentence,
                    "section_target": section,
                    "priority_score": 5,
                }
            elif section in {"methods", "results"} and len(words) >= 34 and sentence.count(",") >= 3:
                rewrite = _rewrite_candidate(sentence)
                entry = {
                    "paragraph_index": pidx,
                    "issue_type": "clarity",
                    "severity": "medium",
                    "critique": f'This sentence compresses too many procedural details into one unit: "{sentence[:220]}". Split the action from the purpose or readout.',
                    "suggested_revision": f"Proposed edit: {rewrite}",
                    "rationale": "Sentence-level methods/results clarity issue grounded in local text.",
                    "locked_paragraph": True,
                    "anchor_text": sentence,
                    "span_sentence": sentence,
                    "section_target": section,
                    "priority_score": 4,
                }
            elif section == "introduction" and len(words) >= 28 and any(
                phrase in slow for phrase in ["to bridge this gap", "for these reasons", "recently", "however,"]
            ):
                rewrite = _rewrite_candidate(sentence)
                entry = {
                    "paragraph_index": pidx,
                    "issue_type": "clarity",
                    "severity": "medium",
                    "critique": f'This framing sentence is doing too much at once: "{sentence[:220]}". Tighten the gap statement so the problem and contribution are easier to follow.',
                    "suggested_revision": f"Proposed edit: {rewrite}",
                    "rationale": "Sentence-level framing cleanup grounded in introduction text.",
                    "locked_paragraph": True,
                    "anchor_text": sentence,
                    "span_sentence": sentence,
                    "section_target": section,
                    "priority_score": 3,
                }
            elif section in {"discussion", "conclusions"} and len(words) >= 20 and any(
                phrase in slow for phrase in ["we conclude", "this study", "this work", "showcases", "impact"]
            ):
                rewrite = _calibrate_claim_sentence(sentence)
                if rewrite != sentence:
                    entry = {
                        "paragraph_index": pidx,
                        "issue_type": "evidence/overclaim concern",
                        "severity": "high",
                        "critique": f'This interpretation sentence would be stronger with narrower scope language: "{sentence[:220]}".',
                        "suggested_revision": f"Proposed edit: {rewrite}",
                        "rationale": "Sentence-level conclusion/discussion calibration grounded in local text.",
                        "locked_paragraph": True,
                        "anchor_text": sentence,
                        "span_sentence": sentence,
                        "section_target": section,
                        "priority_score": 4,
                    }
            if entry:
                sec = str(entry.get("section_target", "body"))
                if per_section.get(sec, 0) >= 4:
                    continue
                candidates.append(entry)
                per_section[sec] = per_section.get(sec, 0) + 1
                break
    return candidates


def review_to_comment_entries(
    review: Any,
    doc: ParsedDocument | None = None,
    base_docx: Path | None = None,
    max_comments: int = 36,
) -> list[dict[str, Any]]:
    entries: list[dict[str, Any]] = []
    if base_docx is not None:
        entries.extend(_build_sentence_level_candidates(base_docx, max_comments=max_comments))
    for idx, sec in enumerate(review.section_specific_comments, start=1):
        critique = str(sec.comment).strip()
        if not critique:
            continue
        if "add one concrete evidence statement" in critique.lower():
            continue
        entries.append(
            {
                "paragraph_index": idx,
                "issue_type": "section_issue",
                "severity": sec.severity,
                "critique": critique,
                "suggested_revision": "Proposed edit: revise the sentence that carries this claim so it states one concrete condition and one explicit limitation.",
                "rationale": "Section-specific issue from review output with local revision guidance.",
                "section_target": _normalize_section_target(getattr(sec, "section", "")),
                "priority_score": 2,
            }
        )
    for idx, item in enumerate(review.extracted_action_items, start=len(entries) + 1):
        action = str(item.action).strip()
        if not action:
            continue
        if action.lower().startswith("address:"):
            action = action.split(":", 1)[-1].strip()
        entries.append(
            {
                "paragraph_index": idx,
                "issue_type": "clarity",
                "severity": item.priority,
                "critique": action,
                "suggested_revision": (
                    "Suggested rewrite: replace vague phrasing with one sentence naming the condition/metric/outcome and one sentence stating boundary conditions."
                ),
                "rationale": "Derived from extracted action item requiring manuscript-local clarification.",
                "priority_score": 1,
            }
        )
        if len(entries) >= max_comments:
            break
    category_pairs = [
        ("methods concern", getattr(review, "methodological_concerns", [])),
        ("evidence/overclaim concern", getattr(review, "novelty_concerns", [])),
        ("citation/reference concern", getattr(review, "citation_reference_concerns", [])),
        ("grammar/style", getattr(review, "writing_organization_concerns", [])),
        ("formatting/journal style", getattr(review, "figure_table_concerns", [])),
        ("redundancy", getattr(review, "major_weaknesses", [])),
    ]
    for label, items in category_pairs:
        for item in items[:4]:
            text = str(item).strip()
            if not text:
                continue
            if label == "methods concern":
                suggestion = (
                    "Add a sentence that names the control/baseline explicitly and links it to the exact experiment or metric."
                )
            elif label == "evidence/overclaim concern":
                suggestion = (
                    "Tone down certainty language and add a boundary-condition sentence that states what was not tested."
                )
            elif label == "citation/reference concern":
                suggestion = (
                    "Attach a specific citation and state whether the support is direct evidence or contextual prior work."
                )
            elif label == "grammar/style":
                suggestion = "Shorten the sentence, prefer active voice, and keep one claim per sentence."
            elif label == "formatting/journal style":
                suggestion = "Align figure/table callout wording and caption style with journal and lab formatting guidance."
            else:
                suggestion = "Remove repetition and merge overlapping statements into one precise sentence."
            entries.append(
                {
                    "paragraph_index": len(entries) + 1,
                    "issue_type": label,
                    "severity": "medium",
                    "critique": text,
                    "suggested_revision": suggestion,
                    "rationale": f"Derived from {label} findings.",
                    "priority_score": 1,
                }
            )
            if len(entries) >= max_comments:
                break
        if len(entries) >= max_comments:
            break
    detail_candidates = [str(t).strip() for t in getattr(review, "detailed_reviewer_comments", []) if str(t).strip()]
    for text in detail_candidates:
        if len(entries) >= max_comments:
            break
        entries.append(
            {
                "paragraph_index": len(entries) + 1,
                "issue_type": "clarity",
                "severity": "medium",
                "critique": text,
                "suggested_revision": "Suggested rewrite: tighten the sentence to one core claim, one supporting detail, and remove ambiguous qualifiers.",
                "rationale": "Detailed reviewer comment promoted into local actionable guidance.",
                "priority_score": 1,
            }
        )
    if not entries:
        for idx, text in enumerate(review.detailed_reviewer_comments[:max_comments], start=1):
            entries.append(
                {
                    "paragraph_index": idx,
                    "issue_type": "grammar/style",
                    "severity": "medium",
                    "critique": str(text).strip(),
                    "suggested_revision": "Rewrite this sentence with tighter wording and explicit subject-action-result structure.",
                    "rationale": "Fallback from detailed reviewer comments with sentence-level edit intent.",
                }
            )
    if not entries:
        summary = str(getattr(review, "summary", "") or "").strip()
        weaknesses = [str(x) for x in getattr(review, "major_weaknesses", []) if str(x).strip()]
        seed = weaknesses[:3] or ([summary] if summary else [])
        if not seed:
            seed = [
                "Clarify core contribution and novelty claim.",
                "Tighten methods/result linkage language.",
                "Improve writing precision and remove ambiguous phrasing.",
            ]
        for idx, text in enumerate(seed, start=1):
            entries.append(
                {
                    "paragraph_index": idx,
                    "issue_type": "clarity",
                    "severity": "medium",
                    "critique": text,
                    "suggested_revision": "Replace vague phrasing with one specific claim and one concrete supporting detail.",
                    "rationale": "Fallback comment when structured review fields were sparse.",
                }
            )

    # Inject sentence-level edits from manuscript text for better coverage and utility.
    if base_docx is not None:
        paragraphs = _analysis_paragraphs(base_docx)
        per_para_added: dict[int, int] = {}
        for pidx, text in enumerate(paragraphs):
            if len(entries) >= max_comments:
                break
            if not text or len(text.split()) < 18:
                continue
            low = text.lower()
            if re.match(r"^abstract\b", low) or low.startswith("abstract:"):
                continue
            if re.fullmatch(r"\d{3,5}", text.strip()):
                continue
            if _is_placeholder_paragraph(low):
                continue
            if _is_header_footer_text(low):
                continue
            if re.match(r"^#+\s", text.strip()) or re.match(r"^■", text.strip()):
                continue
            if _is_front_or_back_matter_text(low):
                continue
            if "abstract" in low and len(text.split()) < 30:
                continue
            sentences = [s.strip() for s in re.split(r"(?<=[.!?])\s+", text) if s.strip()]
            for s in sentences:
                if len(entries) >= max_comments:
                    break
                if len(s.split()) < 14:
                    continue
                if per_para_added.get(pidx, 0) >= 2:
                    break
                slow = s.lower()
                issue_type = None
                critique = None
                rewrite = None
                if " was " in f" {slow} " or " were " in f" {slow} ":
                    rewrite = _rewrite_candidate(s)
                    issue_type = "grammar/style"
                    critique = f"Passive construction weakens readability and agency: {s[:220]}"
                elif len(s.split()) >= 34:
                    rewrite = _rewrite_candidate(s)
                    issue_type = "clarity"
                    critique = f"Overlong sentence likely burdens comprehension: {s[:220]}"
                elif any(k in slow for k in ["very", "highly", "robust", "always", "never", "every", "all "]):
                    rewrite = _rewrite_candidate(s)
                    issue_type = "evidence/overclaim concern"
                    critique = f"Potentially overstated wording should be calibrated to evidence scope: {s[:220]}"
                elif any(k in slow for k in ["this suggests", "this indicates", "this demonstrates"]) and len(s.split()) > 18:
                    rewrite = _rewrite_candidate(s)
                    issue_type = "clarity"
                    critique = f"Inference phrase may need tighter grounding and explicit condition: {s[:220]}"
                if issue_type and critique and rewrite:
                    entries.append(
                        {
                            "paragraph_index": pidx,
                            "issue_type": issue_type,
                            "severity": "medium",
                            "critique": critique,
                            "suggested_revision": f"Suggested rewrite: {rewrite}",
                            "rationale": "Sentence-level editorial pass grounded in local manuscript text.",
                            "locked_paragraph": True,
                            "anchor_text": s,
                            "span_sentence": s,
                            "section_target": None,
                            "priority_score": 3,
                        }
                    )
                    per_para_added[pidx] = per_para_added.get(pidx, 0) + 1

    # Add manuscript-specific claim calibration prompts if phrases appear.
    if doc is not None:
        text = doc.cleaned_text.lower()
        claim_hits = []
        for phrase in ["first attempt", "every instance tried", "modest to excellent yields"]:
            if phrase in text:
                claim_hits.append(phrase)
        for phrase in claim_hits[:3]:
            anchor_index = None
            if base_docx is not None:
                for i, text in enumerate(_analysis_paragraphs(base_docx)):
                    if phrase in text.lower():
                        anchor_index = i
                        break
            entry = {
                "paragraph_index": anchor_index if anchor_index is not None else len(entries) + 1,
                "issue_type": "evidence/overclaim concern",
                "severity": "high",
                "critique": f"Claim language around '{phrase}' may overstate generality without clarifying assay vs isolated yields or scope.",
                "suggested_revision": (
                    "Clarify whether this refers to at least one successful condition per case study, "
                    "distinguish assay yield vs isolated yield, and specify the scope of 'first attempt'."
                ),
                "rationale": "Manuscript-specific claim calibration cue detected in text.",
                "allow_abstract": True,
                "anchor_phrase": phrase,
                "anchor_text": phrase,
                "priority_score": 5,
            }
            if anchor_index is not None:
                entry["locked_paragraph"] = True
            entries.append(entry)
    clean: list[dict[str, Any]] = []
    seen: set[str] = set()
    for e in entries:
        if not e.get("comment_id"):
            e["comment_id"] = f"cmt_{uuid4().hex[:10]}"
        critique = re.sub(r"\s+", " ", str(e.get("critique", "")).strip())
        suggestion = re.sub(r"\s+", " ", str(e.get("suggested_revision", "")).strip())
        if not critique:
            continue
        if any(x in suggestion.lower() for x in ["apply action:", "address:", "revise wording to address this issue"]):
            suggestion = "Rewrite this location with concrete condition, evidence statement, and limitation language."
            e["suggested_revision"] = suggestion
        if _looks_generic_comment(critique, suggestion):
            continue
        if _is_absurd_comment(critique, suggestion):
            continue
        e["critique"] = critique
        e["suggested_revision"] = suggestion
        sig = (e.get("issue_type", ""), critique[:180].lower(), str(e.get("paragraph_index")))
        if sig in seen:
            continue
        seen.add(sig)
        clean.append(e)
        if len(clean) >= max_comments:
            break
    if not clean:
        fallback_text = ""
        detail_comments = [str(x).strip() for x in getattr(review, "detailed_reviewer_comments", []) if str(x).strip()]
        if detail_comments:
            fallback_text = detail_comments[0]
        else:
            fallback_text = "Tighten this paragraph by replacing vague wording with one concrete claim and one supporting detail."
        clean.append(
            {
                "comment_id": f"cmt_{uuid4().hex[:10]}",
                "paragraph_index": 0,
                "issue_type": "clarity",
                "severity": "medium",
                "critique": fallback_text,
                "suggested_revision": "Proposed edit: rewrite the first sentence to clearly state the claim and evidence.",
                "rationale": "Fallback to ensure at least one actionable comment when all generated comments are filtered.",
            }
        )
    return clean


def _rewrite_candidate(sentence: str) -> str:
    s = sentence.strip()
    s = re.sub(r"\s+", " ", s)
    s = re.sub(r"\s+([,.;:])", r"\1", s)
    s = re.sub(r"\[\s*[,;]\s*\]", "", s)
    if len(s) > 260:
        s = s[:260].rstrip() + "..."
    clauses = [c.strip() for c in s.split(",") if c.strip()]
    if s.lower().startswith("to ") and len(clauses) >= 2:
        words = s.split()
        short = " ".join(words[:26]).rstrip(".")
        return short + "."
    if len(clauses) >= 2:
        c1 = clauses[0].rstrip(".")
        c2 = clauses[1][0:120].rstrip(".")
        if len(c1.split()) >= 6 and len(c2.split()) >= 6:
            if c2 and c2[0].islower():
                c2 = c2[0].upper() + c2[1:]
            return f"{c1}. {c2}."
    words = s.split()
    if len(words) > 24:
        short = " ".join(words[:24]).rstrip(".")
        return short + "."
    return s if s.endswith(".") else s + "."


_COMMENT_STOPWORDS = {
    "this", "that", "these", "those", "with", "from", "into", "onto", "over", "under", "between",
    "while", "where", "which", "their", "there", "here", "have", "been", "being", "were", "was",
    "also", "than", "then", "them", "they", "does", "doesnt", "should", "could", "would", "using",
    "used", "use", "such", "more", "less", "many", "much", "some", "often", "very", "however",
}


def _meaningful_tokens(text: str, min_len: int = 4) -> list[str]:
    cleaned = re.sub(r"\[[^\]]+\]", " ", text.lower())
    tokens = re.findall(r"[a-z0-9][a-z0-9\-_/]+", cleaned)
    out: list[str] = []
    for tok in tokens:
        if len(tok) < min_len:
            continue
        if tok in _COMMENT_STOPWORDS:
            continue
        out.append(tok)
    return out


def _short_quote(text: str, max_words: int = 24) -> str:
    words = text.strip().split()
    if len(words) <= max_words:
        return text.strip()
    return " ".join(words[:max_words]).rstrip(" ,;:") + " ..."


def _infer_issue_type_group(issue_type: str) -> str:
    low = issue_type.lower().strip()
    if low in {"evidence/overclaim concern", "novelty concern"}:
        return "evidence"
    if low in {"methods concern", "reproducibility", "statistics concern"}:
        return "methods"
    if low in {"citation/reference concern"}:
        return "citation"
    if low in {"redundancy"}:
        return "redundancy"
    if low in {"grammar/style"}:
        return "style"
    if low in {"structure/organization", "section_issue"}:
        return "structure"
    return "clarity"


def _comment_style_for_section(section: str, issue_group: str, target: str) -> tuple[str, str]:
    quote = _short_quote(target)
    if issue_group == "evidence":
        critique = (
            f'This sentence currently reads broader than the evidence it names: "{quote}". '
            "State the tested case, condition, or measured outcome directly so the reader can see the boundary of the claim."
        )
        suggestion = (
            "Suggested wording direction: qualify the claim to the tested set or experiment, and name the concrete readout "
            "instead of implying general success."
        )
        return critique, suggestion
    if issue_group == "methods":
        critique = (
            f'This procedural sentence bundles setup details without making the decision point clear: "{quote}". '
            "Separate the operation from the criterion, control, or readout that tells the reader why this step matters."
        )
        suggestion = (
            "Suggested wording direction: keep the operation in one sentence, then add a second sentence naming the parameter, "
            "control, or readout that governs interpretation."
        )
        return critique, suggestion
    if issue_group == "citation":
        critique = (
            f'This sentence makes a factual or comparative claim without making the support relationship explicit: "{quote}". '
            "Either point to the exact supporting citation here or qualify the statement as background rather than direct evidence."
        )
        suggestion = (
            "Suggested wording direction: attach the precise citation at the claim site, or soften the sentence so it reads as context "
            "instead of verified support."
        )
        return critique, suggestion
    if issue_group == "redundancy":
        critique = (
            f'This sentence repeats setup or context that nearby text already covers: "{quote}". '
            "Keep the unique point here and cut the repeated setup language."
        )
        suggestion = (
            "Suggested wording direction: retain the new claim or decision in this sentence and delete the repeated setup or motivation phrase."
        )
        return critique, suggestion
    if issue_group == "style":
        critique = (
            f'This sentence is readable but still heavy at the clause level: "{quote}". '
            "Tighten the subject-action-result sequence so the reader does not have to parse multiple weak turns."
        )
        suggestion = (
            "Suggested wording direction: move the main actor and verb earlier, trim filler modifiers, and keep one core action per sentence."
        )
        return critique, suggestion
    if section == "introduction":
        critique = (
            f'This sentence is carrying both background and the paper-specific turn: "{quote}". '
            "Separate the field context from this manuscript's contribution so the reader can see the transition sooner."
        )
        suggestion = (
            "Suggested wording direction: keep the field-level context in one sentence and move this paper's contribution or gap statement into the next."
        )
        return critique, suggestion
    if section == "methods":
        critique = (
            f'This methods sentence names the action but not the exact interpretive boundary the reader needs: "{quote}". '
            "State the condition, criterion, or scope limit at the same point as the step."
        )
        suggestion = (
            "Suggested wording direction: name the operation first, then add the exact condition, threshold, or scope boundary that controls the step."
        )
        return critique, suggestion
    if section == "results":
        critique = (
            f'This results sentence mixes the outcome with too much setup detail: "{quote}". '
            "Lead with the observed result, then move conditions or reagent detail to a follow-on sentence."
        )
        suggestion = (
            "Suggested wording direction: state the observed outcome first, then place the condition, reagent list, or comparison in a second sentence."
        )
        return critique, suggestion
    if section in {"discussion", "conclusions"}:
        critique = (
            f'This interpretation sentence would be stronger with a sharper scope boundary: "{quote}". '
            "Make the takeaway explicit, but keep it tied to the evidence actually shown in the paper."
        )
        suggestion = (
            "Suggested wording direction: keep the interpretation, but add the limiting condition or evidence boundary in the same sentence."
        )
        return critique, suggestion
    critique = (
        f'This sentence is not yet doing one clear editorial job: "{quote}". '
        "Clarify the main point and remove the extra turn that forces the reader to infer the sentence's real function."
    )
    suggestion = (
        "Suggested wording direction: keep one main point in this sentence and move any supporting condition, exception, or qualification to a follow-on sentence."
    )
    return critique, suggestion


def _comment_rationale(section: str, issue_group: str, target: str, paragraph: str) -> str:
    cue_tokens = _meaningful_tokens(target)[:3] or _meaningful_tokens(paragraph)[:3]
    cue_text = ", ".join(cue_tokens) if cue_tokens else "local manuscript wording"
    return (
        f"Anchored to the quoted sentence in the {section or 'body'} section using local cues ({cue_text}); "
        f"the issue is treated as a sentence-level {issue_group} problem rather than a whole-section rewrite request."
    )


def _calibrate_comment_severity(issue_group: str, target: str) -> str:
    low = target.lower()
    if issue_group == "evidence":
        return "high"
    if issue_group == "citation":
        return "medium"
    if issue_group == "methods" and any(k in low for k in ["control", "benchmark", "compared", "versus", "criteria", "criterion"]):
        return "high"
    if issue_group == "redundancy":
        return "low"
    return "medium"


def _is_absurd_comment(critique: str, suggestion: str) -> bool:
    text = f"{critique} {suggestion}".lower()
    blocked_patterns = [
        "remove author",
        "delete author",
        "drop author",
        "remove corresponding author",
        "change corresponding author",
        "remove pi",
        "replace author list",
        "delete affiliation",
    ]
    return any(p in text for p in blocked_patterns)


def _looks_generic_comment(critique: str, suggestion: str) -> bool:
    c = critique.lower().strip()
    s = suggestion.lower().strip()
    if len(c) < 28:
        return True
    generic_starts = ("clarify", "improve", "discuss", "elaborate", "add detail", "address this", "revise this", "rewrite this")
    if c.startswith(generic_starts):
        return True
    if s in {"", "suggested rewrite:", "proposed edit:", "suggested wording direction:"}:
        return True
    if "improve clarity" in s and len(s.split()) < 8:
        return True
    filler_phrases = [
        "section could benefit",
        "needs significant expansion",
        "provide more detail",
        "add more detail",
        "lacks sufficient",
        "more information",
        "helpful to include",
    ]
    if any(p in c for p in filler_phrases):
        return True
    return False


def _comment_entry_quality_ok(entry: dict[str, Any], paragraph_text: str) -> bool:
    critique = str(entry.get("critique", "")).strip()
    suggestion = str(entry.get("suggested_revision", "")).strip()
    anchor = str(entry.get("anchor_text", "") or entry.get("span_sentence", "")).strip()
    if not critique:
        return False
    if _looks_generic_comment(critique, suggestion):
        return False
    if suggestion and len([w for w in re.split(r"\W+", suggestion) if w]) < 8:
        return False
    if anchor and len(anchor.split()) < 6:
        return False
    if paragraph_text and anchor and anchor not in paragraph_text:
        return False
    if paragraph_text:
        linked = _token_overlap(paragraph_text, f"{critique} {suggestion} {anchor}") >= 0.06
        anchor_linked = _token_overlap(anchor or paragraph_text, critique) >= 0.08
        if not linked and not anchor_linked:
            return False
        critique_low = critique.lower()
        if anchor and len(_meaningful_tokens(anchor)) >= 2:
            shared = set(_meaningful_tokens(anchor)) & set(_meaningful_tokens(critique))
            if len(shared) < 1 and not any(k in critique_low for k in ["scope", "boundary", "control", "citation", "outcome", "readout", "sentence"]):
                return False
    return True


def _revise_comment_entries(entries: list[dict[str, Any]], base_docx: Path) -> list[dict[str, Any]]:
    paragraphs = _analysis_paragraphs(base_docx)
    section_by_idx = _build_section_index_map(paragraphs)
    revised: list[dict[str, Any]] = []
    for entry in entries:
        pidx = entry.get("paragraph_index")
        if not isinstance(pidx, int) or not (0 <= pidx < len(paragraphs)):
            revised.append(entry)
            continue
        paragraph = paragraphs[pidx]
        section = section_by_idx.get(pidx, "body")
        target = str(entry.get("anchor_text", "") or entry.get("span_sentence", "")).strip()
        if not target:
            target = _split_sentences(paragraph)[0] if _split_sentences(paragraph) else paragraph[:240]
            entry["anchor_text"] = target
            entry["span_sentence"] = target
        issue_group = _infer_issue_type_group(str(entry.get("issue_type", "")))
        critique, suggestion = _comment_style_for_section(section, issue_group, target)
        entry["critique"] = critique
        entry["suggested_revision"] = suggestion
        entry["rationale"] = _comment_rationale(section, issue_group, target, paragraph)
        entry["severity"] = _calibrate_comment_severity(issue_group, target)
        revised.append(entry)
    return revised


def _dedupe_comment_entries(entries: list[dict[str, Any]]) -> list[dict[str, Any]]:
    if not entries:
        return entries
    severity_rank = {"high": 3, "medium": 2, "low": 1}
    sorted_entries = sorted(
        entries,
        key=lambda e: (
            -severity_rank.get(str(e.get("severity", "medium")).lower(), 2),
            -(int(e.get("priority_score", 0))),
        ),
    )
    kept: list[dict[str, Any]] = []
    seen: set[tuple[str, int, str]] = set()
    for entry in sorted_entries:
        pidx = entry.get("paragraph_index")
        if not isinstance(pidx, int):
            continue
        issue_group = _infer_issue_type_group(str(entry.get("issue_type", "")))
        anchor = str(entry.get("anchor_text", "") or entry.get("span_sentence", "")).strip().lower()
        anchor_sig = " ".join(_meaningful_tokens(anchor, min_len=3)[:8]) or anchor[:120]
        sig = (issue_group, pidx, anchor_sig)
        if sig in seen:
            continue
        seen.add(sig)
        kept.append(entry)
    kept_ids = {id(e) for e in kept}
    return [e for e in entries if id(e) in kept_ids]


def _filter_comment_entries_by_paragraph_quality(entries: list[dict[str, Any]], base_docx: Path) -> list[dict[str, Any]]:
    paragraphs = _analysis_paragraphs(base_docx)
    filtered: list[dict[str, Any]] = []
    for e in entries:
        pidx = e.get("paragraph_index")
        para = ""
        if isinstance(pidx, int) and 0 <= pidx < len(paragraphs):
            para = paragraphs[pidx]
        if _comment_entry_quality_ok(e, para):
            filtered.append(e)
    return filtered


def _limit_comments_per_paragraph(entries: list[dict[str, Any]], max_per_paragraph: int = 2) -> list[dict[str, Any]]:
    if not entries:
        return entries
    severity_rank = {"high": 0, "medium": 1, "low": 2}
    type_rank = {
        "evidence/overclaim concern": 0,
        "methods concern": 1,
        "clarity": 2,
        "grammar/style": 3,
        "redundancy": 4,
        "structure/organization": 5,
        "citation/reference concern": 6,
        "formatting/journal style": 7,
        "figure/table concern": 8,
    }
    grouped: dict[int, list[dict[str, Any]]] = {}
    for e in entries:
        pidx = e.get("paragraph_index")
        if not isinstance(pidx, int):
            continue
        grouped.setdefault(pidx, []).append(e)
    trimmed: list[dict[str, Any]] = []
    used_ids = set()
    for pidx, group in grouped.items():
        if len(group) <= max_per_paragraph:
            trimmed.extend(group)
            continue
        # Prefer highest severity and diverse issue types.
        group_sorted = sorted(
            group,
            key=lambda x: (
                type_rank.get(str(x.get("issue_type", "")).lower(), 9),
                severity_rank.get(str(x.get("severity", "medium")).lower(), 1),
            ),
        )
        kept: list[dict[str, Any]] = []
        seen_types: set[str] = set()
        for item in group_sorted:
            itype = str(item.get("issue_type", "")).lower()
            if itype in seen_types and itype == "evidence/overclaim concern":
                continue
            kept.append(item)
            seen_types.add(itype)
            if len(kept) >= max_per_paragraph:
                break
        trimmed.extend(kept)
    # Preserve original order for stability.
    for e in entries:
        if id(e) in used_ids:
            continue
    kept_set = {id(e) for e in trimmed}
    return [e for e in entries if id(e) in kept_set]


def _paper_context_snapshot(base_docx: Path, max_chars: int = 900) -> str:
    paragraphs = _analysis_paragraphs(base_docx)
    section_by_idx = _build_section_index_map(paragraphs)
    sections = ["abstract", "introduction", "methods", "results", "discussion", "conclusions"]
    lines: list[str] = []
    for section in sections:
        seen = 0
        for idx, text in enumerate(paragraphs):
            if section_by_idx.get(idx) != section:
                continue
            cleaned = text.strip()
            if not cleaned or _is_heading_paragraph(cleaned):
                continue
            lines.append(f"[{section}] {cleaned[:180]}")
            seen += 1
            if seen >= 2:
                break
    snapshot = "\n".join(lines).strip()
    return snapshot[:max_chars]


def _final_comment_arbitration(
    entries: list[dict[str, Any]],
    base_docx: Path,
    provider: Provider | None,
    model: str | None,
    timeout_seconds: int,
) -> list[dict[str, Any]]:
    if not entries or provider is None or not model:
        return entries
    paragraphs = _analysis_paragraphs(base_docx)
    section_by_idx = _build_section_index_map(paragraphs)
    paper_snapshot = _paper_context_snapshot(base_docx)
    items: list[dict[str, Any]] = []
    index_by_id: dict[str, dict[str, Any]] = {}
    for entry in entries:
        pidx = entry.get("paragraph_index")
        if not isinstance(pidx, int) or not (0 <= pidx < len(paragraphs)):
            continue
        paragraph = paragraphs[pidx]
        section = section_by_idx.get(pidx, "body")
        anchor = str(entry.get("anchor_text", "") or entry.get("span_sentence", "")).strip()
        cid = str(entry.get("comment_id") or f"cmt_{uuid4().hex[:10]}")
        entry["comment_id"] = cid
        index_by_id[cid] = entry
        items.append(
            {
                "comment_id": cid,
                "section": section,
                "anchor_sentence": anchor[:320],
                "paragraph": paragraph[:420],
                "previous_paragraph": paragraphs[pidx - 1][:180] if pidx > 0 else "",
                "next_paragraph": paragraphs[pidx + 1][:180] if pidx + 1 < len(paragraphs) else "",
                "candidate_comment": {
                    "issue_type": str(entry.get("issue_type", "")),
                    "severity": str(entry.get("severity", "medium")),
                    "critique": str(entry.get("critique", ""))[:420],
                    "suggested_revision": str(entry.get("suggested_revision", ""))[:420],
                    "rationale": str(entry.get("rationale", ""))[:220],
                },
            }
        )
    if not items:
        return entries
    payload = {
        "paper_snapshot": paper_snapshot[:700],
        "items": items,
    }
    system_prompt = (
        "You are the final editorial arbiter for inline manuscript comments. "
        "Judge each candidate comment in context. Return ONLY JSON with key decisions, where decisions is a list of objects with keys: "
        "comment_id, action, issue_type, severity, critique, suggested_revision, rationale, deletion_reason. "
        "action must be keep, revise, or drop. Drop comments that are generic, weakly grounded, redundant, or not author-helpful."
    )
    try:
        resp = provider.chat(
            ChatRequest(
                model=model,
                system_prompt=system_prompt,
                user_prompt=json.dumps(payload, ensure_ascii=False),
                temperature=0.0,
                max_tokens=1400,
                timeout_seconds=timeout_seconds,
                metadata={"purpose": "final_comment_arbitration_batch"},
            )
        )
        parsed = json.loads(extract_json_candidate(resp.content) or resp.content)
    except Exception:
        parsed = {"decisions": []}
    decisions = parsed.get("decisions", []) if isinstance(parsed, dict) else []
    if not decisions and len(items) == 1 and isinstance(parsed, dict):
        single = dict(parsed)
        single.setdefault("comment_id", items[0]["comment_id"])
        decisions = [single]
    decisions_by_id = {
        str(item.get("comment_id")): item
        for item in decisions
        if isinstance(item, dict) and str(item.get("comment_id", "")).strip()
    }
    revised_entries: list[dict[str, Any]] = []
    for entry in entries:
        cid = str(entry.get("comment_id", ""))
        if cid not in index_by_id:
            revised_entries.append(entry)
            continue
        paragraph = paragraphs[int(entry["paragraph_index"])]
        parsed = decisions_by_id.get(cid, {"action": "keep"})
        action = str(parsed.get("action", "keep")).strip().lower()
        if action == "drop":
            continue
        candidate = dict(entry)
        if action == "revise":
            for key in ["issue_type", "severity", "critique", "suggested_revision", "rationale"]:
                value = parsed.get(key, candidate.get(key))
                if isinstance(value, str) and value.strip():
                    candidate[key] = value.strip()
        if _comment_entry_quality_ok(candidate, paragraph):
            revised_entries.append(candidate)
        elif _comment_entry_quality_ok(entry, paragraph):
            revised_entries.append(entry)
    return revised_entries


def _is_front_or_back_matter_text(text: str) -> bool:
    t = text.lower().strip()
    if not t:
        return True
    if re.fullmatch(r"[a-z0-9][a-z0-9_.-]{10,}", t) and " " not in t:
        # Common filename / slug bleed from managed materials or PDF extraction.
        return True
    if len(t) < 180 and t.count(",") >= 2 and "." not in t:
        # Likely author list / affiliation roster line.
        return True
    if re.match(r"^>?\s*\d+\s*[a-z].*\b(university|department|present address|e-mail:|usa|uk|germany|france|spain|switzerland)\b", t):
        return True
    if re.match(r"^\d+\s*[a-z].*\b(university|department|inc\.|corp\.|usa|uk|germany|france|switzerland)\b", t):
        return True
    if "present address" in t:
        return True
    if "e-mail:" in t or "email:" in t:
        return True
    if re.search(r"\b(received|accepted|published):\b", t):
        return True
    if re.match(r"^\*\*?keywords\b", t) or t.startswith("keywords:"):
        return True
    blocked = [
        "cite this",
        "associated content",
        "supporting information",
        "author information",
        "author contributions",
        "funding",
        "acknowledgment",
        "abbreviations",
        "references",
        "received",
        "accepted",
        "check for updates",
        "nature synthesis",
        "doi.org",
        "acs",
        "keywords",
        "corresponding author",
        "notes",
        "american chemical society",
        "org. process res. dev",
        "organic process research",
        "pubs.acs.org",
        "copyright",
        "©",
    ]
    return any(b in t for b in blocked)


def _is_placeholder_paragraph(text: str) -> bool:
    t = text.strip().lower()
    if "==> picture" in t or "intentionally omitted" in t:
        return True
    if "picture" in t and "omitted" in t:
        return True
    if "start of picture text" in t or "end of picture text" in t:
        return True
    if t.startswith("figure ") and "intentionally omitted" in t:
        return True
    return False


def _is_heading_paragraph(text: str) -> bool:
    t = text.strip()
    if re.match(r"^#\s", t):
        return True
    if re.match(r"^##\s", t):
        return True
    if re.match(r"^###\s", t):
        return True
    if re.match(r"^■", t):
        return True
    if re.match(r"^\*\*\[.+\]\*\*$", t):
        return True
    if re.match(r"^##\s*\*\*?.+\*\*?$", t):
        return True
    if t.isupper() and 1 <= len(t.split()) <= 6:
        return True
    if re.match(r"^\*\*[A-Za-z].+\*\*$", t) and len(t.split()) <= 8:
        return True
    norm = _normalize_heading_text(t)
    if _heading_to_section(norm) is not None and len(t.split()) <= 10:
        return True
    return False


def _is_header_footer_text(text: str) -> bool:
    t = text.strip().lower()
    if not t:
        return False
    shortish = len(t.split()) <= 14 or len(t) <= 110
    if "cite this" in t:
        return False
    if t.strip("* _") == "article":
        return True
    digits = re.sub(r"[^0-9]", "", t)
    if digits and 3 <= len(digits) <= 5 and len(t) <= len(digits) + 4:
        return True
    if "nature synthesis" in t and shortish:
        return True
    if shortish and "volume" in t and ("issue" in t or "november" in t or re.search(r"\b20\d{2}\b", t)):
        return True
    if "check for updates" in t and shortish:
        return True
    if shortish and re.search(r"\bvolume\s*\d+\b", t) and re.search(r"\b\d{4}\b", t):
        return True
    if shortish and re.search(r"\bvolume\s*\d+\b", t) and "|" in t:
        return True
    header_terms = [
        "pubs.acs.org",
        "org. process res. dev",
        "organic process research",
        "https://doi.org/10.1021",
        "doi.org/10.1021",
        "doi.org/10.1038",
    ]
    if any(term in t for term in header_terms) and shortish:
        return True
    return False


def _is_front_or_back_matter_heading(text: str) -> bool:
    t = re.sub(r"[^a-zA-Z ]", " ", text).lower()
    t = re.sub(r"\s+", " ", t).strip()
    if not t:
        return False
    return bool(
        re.match(
            r"^(author contributions|funding|associated content|author information|abbreviations|references|notes|corresponding author|authors|acknowledg(e)?ments|data availability|additional information|supplementary information|competing interests)\b",
            t,
        )
    )


def _normalize_heading_text(text: str) -> str:
    cleaned = re.sub(r"[*#_]+", " ", text)
    cleaned = re.sub(r"[^a-zA-Z ]", " ", cleaned).lower()
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    tokens = cleaned.split()
    if len(tokens) <= 3 and "experimental" in tokens:
        return "experimental"
    if len(tokens) <= 3 and "methods" in tokens:
        return "methods"
    return cleaned


def _content_section_hint(text: str, current: str) -> str | None:
    low = text.lower()
    scores: dict[str, float] = {
        "abstract": 0.0,
        "introduction": 0.0,
        "methods": 0.0,
        "results": 0.0,
        "discussion": 0.0,
        "conclusions": 0.0,
    }
    if "abstract:" in low:
        return "abstract"

    intro_phrases = [
        "chemical space exploration",
        "drug discovery",
        "primary bottleneck",
        "for these reasons",
        "to bridge this gap",
        "inspired by moore",
        "promise to enable",
        "recently, generative transformers",
        "recently it has become increasingly possible",
        "with this background",
        "one motivator",
        "emerged as a technology",
        "have emerged as reactions of choice",
    ]
    method_phrases = [
        "general procedure",
        "experimental procedure",
        "prepared at the 1",
        "using the liquid-handling robot",
        "inside an n2 atmosphere glovebox",
        "in a nitrogen-filled glovebox",
        "was weighed into a",
        "were dosed with",
        "source plate",
        "reaction plate",
        "centrifuged",
        "quenched",
        "ambient temperature",
        "uplc",
        "rp-hplc",
    ]
    result_phrases = [
        "yield",
        "conversion",
        "heatmap",
        "selected examples",
        "best-performing",
        "product was observed",
        "repeated on",
        "surveying reaction space",
        "we initiated our studies",
        "we next explored",
        "we chose",
        "ideal conditions",
        "all 1,440 reactions",
        "relative conversion",
        "we found that",
    ]
    discussion_phrases = [
        "we sought to",
        "this study",
        "this work",
        "overall",
        "in summary",
        "important that",
        "impact that",
        "it is therefore important",
        "showcases key principles",
        "while we find",
        "it remains the case",
    ]
    for phrase in intro_phrases:
        if phrase in low:
            scores["introduction"] += 2.0
    for phrase in method_phrases:
        if phrase in low:
            scores["methods"] += 2.0
    for phrase in result_phrases:
        if phrase in low:
            scores["results"] += 2.0
    for phrase in discussion_phrases:
        if phrase in low:
            scores["discussion"] += 2.0

    if low.startswith("figure ") or low.startswith("fig. ") or low.startswith("table "):
        scores["results"] += 3.0
    if re.search(r"\b\d+(\.\d+)?\s*(mmol|equiv\.?|µl|ul|nmol|h)\b", low):
        scores["methods"] += 1.5
    if re.search(r"\b(we conclude|in conclusion|ultimately)\b", low):
        scores["conclusions"] += 3.0
    if "however," in low and current == "introduction":
        scores["introduction"] += 0.5
    if current in scores and current not in {"body", "front_matter", "references", "header_footer"}:
        scores[current] += 0.75

    best_section, best_score = max(scores.items(), key=lambda kv: kv[1])
    if best_score < 1.5:
        return None
    if best_section == "methods" and scores["introduction"] >= best_score - 0.5 and current == "introduction":
        return "introduction"
    if best_section == "results" and scores["introduction"] >= best_score and current in {"body", "introduction"}:
        return "introduction"
    if best_section == "discussion" and scores["conclusions"] >= best_score:
        return "conclusions"
    return best_section


def _heading_to_section(norm: str) -> str | None:
    if not norm:
        return None
    if norm.startswith("abstract"):
        return "abstract"
    if norm in {"introduction", "background"}:
        return "introduction"
    if norm in {"methods", "experimental", "materials methods", "materials and methods"}:
        return "methods"
    if norm.startswith("general procedure") or "procedure" in norm:
        return "methods"
    if norm in {"results", "results and discussion", "findings"}:
        return "results"
    if norm in {"discussion", "interpretation"}:
        return "discussion"
    if norm in {"conclusion", "conclusions"}:
        return "conclusions"
    if norm.startswith("references"):
        return "references"
    if _is_front_or_back_matter_heading(norm):
        return "front_matter"
    return None


def _has_strong_front_matter_signal(text: str) -> bool:
    low = text.lower().strip()
    signals = [
        "cite this",
        "supporting information",
        "author contributions",
        "author information",
        "corresponding author",
        "present address",
        "e-mail:",
        "email:",
        "received:",
        "accepted:",
        "published:",
        "data availability",
        "additional information",
        "competing interests",
        "acknowledgements",
        "funding",
        "doi.org/",
        "american chemical society",
    ]
    if any(sig in low for sig in signals):
        return True
    if re.match(r"^>?\s*\d+\s*[a-z].*\b(university|department|present address|usa|uk|spain|switzerland)\b", low):
        return True
    return False


def _build_section_index_map(paragraphs: list[str]) -> dict[int, str]:
    section_by_idx: dict[int, str] = {}
    current = "body"
    total = len(paragraphs)
    method_headings: list[int] = []
    references_heading: int | None = None
    substantive_started = False
    for idx, text in enumerate(paragraphs):
        if not text:
            continue
        low = text.strip().lower()
        if _is_header_footer_text(low):
            section_by_idx[idx] = "header_footer"
            continue
        if _is_placeholder_paragraph(low):
            section_by_idx[idx] = current if current not in {"body", "front_matter"} else "body"
            continue
        if _is_front_or_back_matter_text(low):
            section_by_idx[idx] = "front_matter"
            if not substantive_started or current == "front_matter" or (references_heading is not None and idx >= references_heading):
                current = "front_matter"
            continue
        if _is_heading_paragraph(text):
            norm = _normalize_heading_text(text)
            heading_section = _heading_to_section(norm)
            if heading_section:
                current = heading_section
                substantive_started = heading_section not in {"front_matter", "references"}
                if heading_section in {"methods", "experimental"}:
                    method_headings.append(idx)
                if heading_section == "references":
                    references_heading = idx
            elif not substantive_started and idx < max(8, int(total * 0.08)):
                section_by_idx[idx] = "front_matter"
                continue
        elif re.match(r"^abstract\b", low):
            current = "abstract"
            substantive_started = True
        elif current == "body":
            hinted = _content_section_hint(text, current)
            if hinted:
                current = hinted
                substantive_started = hinted not in {"front_matter", "references"}
        elif current == "front_matter":
            hinted = _content_section_hint(text, "introduction")
            if hinted and not _has_strong_front_matter_signal(text):
                current = hinted
                substantive_started = True
        else:
            hinted = _content_section_hint(text, current)
            if hinted and current not in {"references", "front_matter"}:
                current = hinted
        section_by_idx[idx] = current

    body_indices = [idx for idx, sec in section_by_idx.items() if sec == "body"]
    if body_indices and len(body_indices) / max(1, len(section_by_idx)) > 0.6:
        first_methods = min(method_headings) if method_headings else None
        for idx in body_indices:
            text = paragraphs[idx].strip()
            low = text.lower()
            if re.match(r"^\\d+\\.\\s", low) and re.search(r"\\b(19|20)\\d{2}\\b", low):
                section_by_idx[idx] = "references"
                continue
            if first_methods is not None and idx < first_methods:
                section_by_idx[idx] = "introduction"
                continue
            if idx < max(6, int(total * 0.12)) and "abstract" not in low:
                section_by_idx[idx] = "introduction"
                continue
            if any(k in low for k in ["method", "experimental", "procedure", "reagent", "reaction conditions"]):
                section_by_idx[idx] = "methods"
                continue
            if any(k in low for k in ["result", "yield", "conversion", "screen", "figure", "table", "data"]):
                section_by_idx[idx] = "results"
                continue
            if any(k in low for k in ["discussion", "overall", "in summary", "we conclude", "implications"]):
                section_by_idx[idx] = "discussion"
                continue

    # Reclassify methods paragraphs that clearly read like results/discussion.
    for idx, sec in list(section_by_idx.items()):
        if sec not in {"methods", "experimental"}:
            continue
        text = paragraphs[idx].strip().lower()
        if any(k in text for k in ["result", "yield", "conversion", "screen", "figure", "table", "data"]):
            section_by_idx[idx] = "results"
            continue
        if any(k in text for k in ["discussion", "overall", "in summary", "we conclude", "implications"]):
            section_by_idx[idx] = "discussion"
            continue
    # Repair common PDF-surrogate pathologies where title/introduction/results bleed into body/front_matter.
    first_methods = min(method_headings) if method_headings else None
    first_results = min((idx for idx, sec in section_by_idx.items() if sec == "results"), default=None)
    for idx, sec in list(section_by_idx.items()):
        text = paragraphs[idx].strip()
        low = text.lower()
        if sec == "body":
            hint = _content_section_hint(text, sec)
            if hint:
                section_by_idx[idx] = hint
                sec = hint
        if sec == "body" and idx < max(8, int(total * 0.08)) and (_is_heading_paragraph(text) or len(text.split()) <= 12):
            section_by_idx[idx] = "front_matter"
        if sec == "front_matter" and first_methods is not None and idx < first_methods:
            if (
                len(text.split()) >= 18
                and not _is_front_or_back_matter_heading(text)
                and not _has_strong_front_matter_signal(text)
            ):
                hint = _content_section_hint(text, "introduction") or "introduction"
                section_by_idx[idx] = hint
        if (
            first_methods is not None
            and idx < first_methods
            and section_by_idx.get(idx) in {"methods", "results", "body"}
            and not _has_strong_front_matter_signal(text)
        ):
            intro_hint = _content_section_hint(text, "introduction")
            if intro_hint == "introduction":
                section_by_idx[idx] = "introduction"
        if (
            first_results is not None
            and idx < first_results
            and section_by_idx.get(idx) == "results"
            and not any(k in low for k in ["fig.", "figure ", "table ", "yield", "conversion", "repeated on", "all 1,440 reactions"])
        ):
            intro_hint = _content_section_hint(text, "introduction")
            if intro_hint == "introduction":
                section_by_idx[idx] = "introduction"
        if sec == "discussion" and idx > 0 and section_by_idx.get(idx - 1) == "methods" and len(text.split()) >= 18:
            if any(k in low for k in ["glovebox", "equiv", "mmol", "purified", "stirred", "quenched"]):
                section_by_idx[idx] = "methods"
    return section_by_idx


def _assign_paragraph_indices(entries: list[dict[str, Any]], base_docx: Path) -> list[dict[str, Any]]:
    paragraphs = _analysis_paragraphs(base_docx)
    non_empty = [idx for idx, text in enumerate(paragraphs) if text]
    if not non_empty:
        return entries

    section_by_idx = _build_section_index_map(paragraphs)

    eligible = [
        idx
        for idx in non_empty
        if section_by_idx.get(idx) not in {"front_matter", "references", "header_footer"}
        and not _is_placeholder_paragraph(paragraphs[idx])
        and not _is_heading_paragraph(paragraphs[idx])
        and not _is_header_footer_text(paragraphs[idx])
        and not _is_front_or_back_matter_text(paragraphs[idx])
    ]
    if not eligible:
        eligible = non_empty

    used: dict[int, int] = {}
    for i, entry in enumerate(entries):
        if entry.get("locked_paragraph") and isinstance(entry.get("paragraph_index"), int):
            pidx = int(entry["paragraph_index"])
            if 0 <= pidx < len(paragraphs) and section_by_idx.get(pidx) not in {"front_matter", "references"}:
                section = section_by_idx.get(pidx, "body")
                issue_type = str(entry.get("issue_type", "")).lower()
                if section == "abstract" and not entry.get("allow_abstract", False) and issue_type in {"grammar/style", "clarity", "redundancy"}:
                    continue
                entry["paragraph_excerpt"] = paragraphs[pidx][:180]
                entry["section_hint"] = section
                used[pidx] = used.get(pidx, 0) + 1
                continue
        critique = str(entry.get("critique", "")).lower()
        suggestion = str(entry.get("suggested_revision", "")).lower()
        anchor_phrase = str(entry.get("anchor_phrase", "")).lower().strip()
        section_target = _normalize_section_target(entry.get("section_target"))
        anchor_terms = [w for w in (critique + " " + suggestion).split() if len(w) > 5][:12]
        issue_type = str(entry.get("issue_type", "")).lower()

        chosen: int | None = None
        preferred_sections = []
        if issue_type == "methods concern":
            preferred_sections = ["methods", "experimental"]
        elif issue_type == "evidence/overclaim concern":
            preferred_sections = ["results", "discussion", "conclusions", "abstract"]
        elif issue_type in {"grammar/style", "clarity", "redundancy"}:
            preferred_sections = ["introduction", "methods", "experimental", "results", "discussion"]
        elif issue_type in {"formatting/journal style", "figure/table concern"}:
            preferred_sections = ["results", "discussion"]
        elif issue_type == "structure/organization":
            preferred_sections = ["introduction", "results", "discussion", "conclusions"]

        pool = eligible
        if section_target:
            pool = [p for p in eligible if section_by_idx.get(p, "body") == section_target] or eligible
        elif preferred_sections:
            pool = [p for p in eligible if section_by_idx.get(p, "body") in preferred_sections] or eligible
        if issue_type in {"formatting/journal style", "figure/table concern"}:
            pool = [
                p
                for p in pool
                if "figure" in paragraphs[p].lower()
                and not _is_placeholder_paragraph(paragraphs[p])
            ] or pool
        else:
            pool = [p for p in pool if not _is_placeholder_paragraph(paragraphs[p])]

        if anchor_phrase:
            tokens = [t for t in re.split(r"\s+", anchor_phrase) if len(t) > 2]
            for pidx in pool:
                text = paragraphs[pidx].lower()
                if anchor_phrase in text or (tokens and all(tok in text for tok in tokens)):
                    chosen = pidx
                    break

        for pidx in pool:
            text = paragraphs[pidx].lower()
            section = section_by_idx.get(pidx, "body")
            if section == "abstract" and issue_type in {"methods concern", "reproducibility", "statistics concern"}:
                continue
            if section == "abstract" and not entry.get("allow_abstract", False) and issue_type in {"grammar/style", "clarity", "redundancy"}:
                continue
            if section in {"front_matter", "references", "header_footer"}:
                continue
            if "keywords" in text or text.startswith("keywords:"):
                continue
            if re.fullmatch(r"\d{3,5}", text.strip()):
                continue
            if _is_placeholder_paragraph(text):
                if issue_type not in {"formatting/journal style", "figure/table concern"}:
                    continue
            if "figure" in text and issue_type not in {"formatting/journal style", "figure/table concern"}:
                continue
            if any(term in text for term in anchor_terms):
                if used.get(pidx, 0) < 1:
                    chosen = pidx
                    break

        if chosen is None:
            scored_pool: list[tuple[float, int]] = []
            for pidx in pool:
                text = paragraphs[pidx].lower()
                score = _token_overlap(text, critique + " " + suggestion)
                if section_target and section_by_idx.get(pidx, "body") == section_target:
                    score += 0.25
                score -= used.get(pidx, 0) * 0.15
                scored_pool.append((score, pidx))
            if scored_pool:
                scored_pool.sort(reverse=True)
                if scored_pool[0][0] > 0.04:
                    chosen = scored_pool[0][1]
        if chosen is None:
            dist_total = len(eligible)
            target_rank = max(0, min(dist_total - 1, round((i + 1) * dist_total / (len(entries) + 1))))
            chosen = eligible[target_rank]
            attempts = 0
            while used.get(chosen, 0) >= 1 and attempts < dist_total:
                target_rank = min(dist_total - 1, target_rank + 1)
                chosen = eligible[target_rank]
                attempts += 1

        entry["paragraph_index"] = chosen
        entry["paragraph_excerpt"] = paragraphs[chosen][:180]
        entry["section_hint"] = section_by_idx.get(chosen, "body")
        used[chosen] = used.get(chosen, 0) + 1
    return entries


def _localize_comment_entries(entries: list[dict[str, Any]], base_docx: Path) -> list[dict[str, Any]]:
    paragraphs = _analysis_paragraphs(base_docx)
    section_by_idx = _build_section_index_map(paragraphs)
    for entry in entries:
        pidx = entry.get("paragraph_index")
        if not isinstance(pidx, int) or pidx < 0 or pidx >= len(paragraphs):
            continue
        paragraph = paragraphs[pidx]
        section = section_by_idx.get(pidx, "body")
        sentences = _split_sentences(paragraph)
        if not sentences:
            continue
        critique = str(entry.get("critique", "")).strip()
        suggestion = str(entry.get("suggested_revision", "")).strip()
        issue_type = str(entry.get("issue_type", "")).lower()
        anchor = str(entry.get("anchor_text", "")).strip()
        target = ""
        if anchor and anchor in paragraph:
            target = anchor
        else:
            scored = sorted(
                (( _token_overlap(s, critique + " " + suggestion), s) for s in sentences),
                key=lambda x: x[0],
                reverse=True,
            )
            target = scored[0][1] if scored else sentences[0]
        entry["anchor_text"] = target
        entry["span_sentence"] = target
        target_short = target[:220]
        critique_low = critique.lower()
        if issue_type == "section_issue" or _looks_generic_comment(critique, suggestion):
            if section == "methods":
                entry["critique"] = f'This procedural sentence is clear about the operation but not the decision point or scope boundary: "{target_short}". Add the exact readout, criterion, or limitation that governs this step.'
                entry["suggested_revision"] = f"Proposed edit: {_rewrite_candidate(target)}"
            elif section in {"results", "discussion", "conclusions"}:
                entry["critique"] = f'This sentence states the result broadly without enough local qualification: "{target_short}". Name the exact outcome or narrow the scope here.'
                entry["suggested_revision"] = f"Proposed edit: {_calibrate_claim_sentence(target)}"
            else:
                entry["critique"] = f'This sentence carries the local revision burden for the paragraph: "{target_short}". Tighten the wording so the main point and limitation are easier to follow.'
                entry["suggested_revision"] = f"Proposed edit: {_rewrite_candidate(target)}"
        elif issue_type == "methods concern" and "control" in critique_low:
            entry["critique"] = f'This methods sentence names the operation but not the comparison point the reader needs: "{target_short}". If a control or benchmark is relevant, state it at this point rather than later.'
            entry["suggested_revision"] = f"Proposed edit: {_rewrite_candidate(target)}"
        elif issue_type in {"clarity", "grammar/style"} and "suggested rewrite" not in suggestion.lower():
            entry["suggested_revision"] = f"Proposed edit: {_rewrite_candidate(target)}"
        elif issue_type == "evidence/overclaim concern":
            entry["critique"] = f'This sentence risks reading broader than the tested scope: "{target_short}". Narrow the claim or specify the tested condition directly.'
            entry["suggested_revision"] = f"Proposed edit: {_calibrate_claim_sentence(target)}"
    return entries


def _enrich_comment_suggestions(entries: list[dict[str, Any]], base_docx: Path) -> list[dict[str, Any]]:
    paragraphs = _analysis_paragraphs(base_docx)
    templated_markers = [
        "add a sentence",
        "tone down",
        "attach a specific citation",
        "align figure/table",
        "remove repetition",
        "rewrite this sentence",
        "rewrite this location",
        "suggested rewrite",
        "apply action",
        "revise wording",
    ]
    for entry in entries:
        pidx = entry.get("paragraph_index")
        if not isinstance(pidx, int) or pidx < 0 or pidx >= len(paragraphs):
            continue
        text = paragraphs[pidx]
        if not text:
            continue
        low = text.lower()
        if _is_front_or_back_matter_text(low) or _is_heading_paragraph(text):
            continue
        suggestion = str(entry.get("suggested_revision", "")).strip()
        critique = str(entry.get("critique", "")).strip()
        extracted = None
        for marker in [
            "Rewrite long sentence for clarity:",
            "Long sentence to simplify:",
            "Possible passive construction to rewrite:",
            "Clarify this section’s opening claim:",
            "Clarify this section's opening claim:",
        ]:
            if marker in critique:
                extracted = critique.split(marker, 1)[-1].strip().strip("\"")
                break
        if not extracted and "\"" in critique:
            match = re.search(r"\"(.+?)\"", critique)
            if match:
                extracted = match.group(1).strip()
        if not suggestion:
            suggestion = ""
        if suggestion.lower().startswith("suggested wording direction:"):
            continue
        if extracted:
            rewrite = _rewrite_candidate(extracted)
            entry["suggested_revision"] = f"Proposed edit: {rewrite}"
            entry["rationale"] = "Proposed edit derived from the cited sentence in the critique."
            continue
        if any(marker in suggestion.lower() for marker in templated_markers):
            sentence = re.split(r"(?<=[.!?])\s+", text)[0].strip()
            if sentence:
                rewrite = _rewrite_candidate(sentence)
                entry["suggested_revision"] = f"Proposed edit: {rewrite}"
                entry["rationale"] = "Proposed edit derived from the nearby sentence to improve clarity and flow."
    return entries


def _balance_comment_entries(entries: list[dict[str, Any]], base_docx: Path, max_comments: int) -> list[dict[str, Any]]:
    if not entries:
        return entries
    section_by_idx = _section_lookup_for_docx(base_docx)
    severity_rank = {"high": 3, "medium": 2, "low": 1}
    sorted_entries = sorted(
        entries,
        key=lambda e: (
            -(int(e.get("priority_score", 0))),
            -severity_rank.get(str(e.get("severity", "medium")).lower(), 2),
        ),
    )
    grouped: dict[str, list[dict[str, Any]]] = {}
    for entry in sorted_entries:
        pidx = entry.get("paragraph_index")
        if not isinstance(pidx, int):
            continue
        sec = section_by_idx.get(pidx, "body")
        grouped.setdefault(sec, []).append(entry)

    chosen: list[dict[str, Any]] = []
    used_ids: set[str] = set()
    section_counts: dict[str, int] = {}
    sections = ["abstract", "introduction", "methods", "results", "discussion", "conclusions", "body"]
    section_caps = {
        "abstract": 1,
        "introduction": 3,
        "methods": 4,
        "results": 4,
        "discussion": 3,
        "conclusions": 2,
        "body": 2,
    }
    progressed = True
    while len(chosen) < max_comments and progressed:
        progressed = False
        for section in sections:
            section_entries = grouped.get(section, [])
            used_in_section = section_counts.get(section, 0)
            if used_in_section >= min(section_caps.get(section, 3), len(section_entries)):
                continue
            next_entry = section_entries[used_in_section]
            cid = str(next_entry.get("comment_id", ""))
            if cid in used_ids:
                continue
            chosen.append(next_entry)
            used_ids.add(cid)
            section_counts[section] = used_in_section + 1
            progressed = True
            if len(chosen) >= max_comments:
                break
    chosen_ids = {id(e) for e in chosen}
    return [e for e in entries if id(e) in chosen_ids]


def _section_map_for_docx(base_docx: Path) -> list[dict[str, Any]]:
    paragraphs = _analysis_paragraphs(base_docx)
    out: list[dict[str, Any]] = []
    section_by_idx = _build_section_index_map(paragraphs)
    for idx, text in enumerate(paragraphs):
        if not text:
            continue
        section = section_by_idx.get(idx, "body")
        out.append(
            {
                "paragraph_index": idx,
                "section": section,
                "text_excerpt": text[:200],
            }
        )
    return out


def _section_lookup_for_docx(base_docx: Path) -> dict[int, str]:
    return {item["paragraph_index"]: item["section"] for item in _section_map_for_docx(base_docx)}


def _first_substantive_paragraph_index(base_docx: Path) -> int:
    paragraphs = _analysis_paragraphs(base_docx)
    for idx, text in enumerate(paragraphs):
        cleaned = (text or "").strip()
        if not cleaned:
            continue
        if _is_heading_paragraph(cleaned):
            continue
        low = cleaned.lower()
        if _is_front_or_back_matter_text(low):
            continue
        return idx
    for idx, text in enumerate(paragraphs):
        if (text or "").strip():
            return idx
    return 0


def _severity_rank(value: str | None) -> int:
    low = (value or "").lower()
    if low in {"critical", "high"}:
        return 3
    if low in {"medium", "moderate"}:
        return 2
    return 1


def _token_overlap(a: str, b: str) -> float:
    ta = {t for t in re.split(r"\W+", a.lower()) if t}
    tb = {t for t in re.split(r"\W+", b.lower()) if t}
    if not ta or not tb:
        return 0.0
    return len(ta & tb) / float(len(ta | tb))


def _numeric_tokens(text: str) -> set[str]:
    cleaned = re.sub(r"\[[^\]]+\]", "", text)
    cleaned = cleaned.replace(",", "")
    tokens = re.findall(r"\d+(?:\.\d+)?", cleaned)
    return {t.strip() for t in tokens if t.strip()}


def _content_tokens(text: str) -> set[str]:
    stop = {
        "the", "a", "an", "and", "or", "of", "to", "in", "on", "for", "with", "by", "from", "that", "this", "these",
        "those", "is", "are", "was", "were", "be", "been", "being", "it", "its", "as", "at", "we", "our", "their",
        "can", "may", "might", "could", "should", "would", "into", "than", "then", "also", "however", "therefore",
        "thus", "here", "there", "using", "used", "use", "via", "based", "study", "work", "results", "result",
        "data", "method", "methods",
    }
    return {t for t in re.findall(r"[a-zA-Z][a-zA-Z0-9_-]+", text.lower()) if t not in stop and len(t) > 2}


def _basic_rewrite_checks(original: str, revised: str) -> tuple[bool, str | None]:
    o = original.strip()
    r = revised.strip()
    if not r:
        return False, "empty_revision"
    if r == o:
        return False, "no_change"
    if re.match(r"^\s*#{1,6}\s", r):
        return False, "markdown_heading"
    if re.match(r"^\s*\*\*\[[^\]]+\]\*\*", r):
        return False, "markdown_section_label"
    risky_additions = [
        "comparative studies",
        "we conducted",
        "we compared",
        "control experiments were performed",
        "statistically significant",
        "baseline comparison",
    ]
    o_low = o.lower()
    r_low = r.lower()
    for phrase in risky_additions:
        if phrase in r_low and phrase not in o_low:
            return False, "unsupported_addition"
    if len(r) < 0.5 * len(o) or len(r) > 1.85 * len(o):
        return False, "length_drift"
    orig_nums = _numeric_tokens(o)
    rev_nums = _numeric_tokens(r)
    if orig_nums and not orig_nums.issubset(rev_nums):
        return False, "numeric_loss"
    if _token_overlap(o, r) < 0.35:
        return False, "low_overlap"
    orig_content = _content_tokens(o)
    rev_content = _content_tokens(r)
    novel_content = rev_content - orig_content
    if len(novel_content) >= 8 and len(novel_content) > max(7, int(len(orig_content) * 0.75)):
        return False, "unsupported_addition"
    return True, None


def _basic_span_rewrite_checks(target_span: str, revised_span: str) -> tuple[bool, str | None]:
    target = target_span.strip()
    revised = revised_span.strip()
    if not revised:
        return False, "empty_revision"
    if revised == target:
        return False, "no_change"
    low = revised.lower()
    meta_phrases = (
        "this sentence",
        "this paragraph",
        "the authors should",
        "the manuscript should",
        "the paper should",
        "claim should",
        "reader should",
        "suggested wording direction",
        "proposed edit",
        "suggested rewrite",
    )
    if any(phrase in low for phrase in meta_phrases):
        return False, "editorial_meta_text"
    if re.search(r"[,:;]\s*$", revised):
        return False, "incomplete_phrase"
    if target.endswith((".", "?", "!")) and not revised.endswith((".", "?", "!")):
        return False, "punctuation_incomplete"
    if revised.endswith((" and", " or", " with", " to", " of", " in", " for", " by", " from", " than", " then")):
        return False, "incomplete_phrase"
    if revised.count("(") != revised.count(")"):
        return False, "unbalanced_parentheses"
    if revised.count("[") != revised.count("]"):
        return False, "unbalanced_brackets"
    if len(_split_sentences(target)) == 1 and len(_split_sentences(revised)) > 2:
        return False, "scope_drift"
    if _token_overlap(target, revised) < 0.30:
        return False, "low_span_overlap"
    return True, None


def _resolve_target_span(group: list[dict[str, Any]], original: str) -> str:
    for key in ["span_sentence", "anchor_text"]:
        for item in group:
            text = str(item.get(key, "")).strip()
            if text and text in original:
                for sentence in _split_sentences(original):
                    if text in sentence:
                        return sentence
                return text
    sentences = _split_sentences(original)
    if len(sentences) <= 1:
        return original
    critique = " ".join(f"{c.get('critique', '')} {c.get('suggested_revision', '')}" for c in group)
    scored = sorted((( _token_overlap(s, critique), s) for s in sentences), key=lambda x: x[0], reverse=True)
    if scored and scored[0][0] >= 0.03:
        return scored[0][1]
    return sentences[0]


def _apply_span_rewrite(original: str, target_span: str, revised_span: str) -> str:
    if target_span and target_span in original:
        return original.replace(target_span, revised_span, 1)
    return revised_span


def _group_comments_by_rewrite_target(original: str, group: list[dict[str, Any]]) -> list[tuple[str, list[dict[str, Any]]]]:
    by_target: dict[str, list[dict[str, Any]]] = {}
    for item in group:
        target = _resolve_target_span([item], original).strip() or original.strip()
        by_target.setdefault(target, []).append(item)
    return sorted(by_target.items(), key=lambda pair: original.find(pair[0]) if pair[0] in original else 10**9)


def _requires_nonlocal_addition(group: list[dict[str, Any]], original: str, section: str) -> bool:
    text = " ".join(str(c.get("critique", "")) for c in group).lower()
    if any(k in text for k in ["provide more information", "include more details", "expand the methods section", "lacks sufficient controls", "thorough analysis of uncertainty"]):
        if not any(str(c.get("anchor_text", "")).strip() or str(c.get("span_sentence", "")).strip() for c in group):
            if section in {"methods", "results", "discussion"} and len(original.split()) >= 40:
                return True
    return False


def _is_global_issue_not_localized(group: list[dict[str, Any]], original: str) -> bool:
    text = " ".join([str(c.get("critique", "")) for c in group]).lower()
    if any(k in text for k in ["whole manuscript", "overall structure", "global framing", "throughout the paper"]):
        return True
    # If there is no concrete anchor and only abstract structural language, keep as unresolved.
    has_anchor = any(str(c.get("anchor_text", "")).strip() for c in group)
    structural_only = all(str(c.get("issue_type", "")).lower() in {"structure/organization", "framing", "global"} for c in group)
    if structural_only and not has_anchor and not re.search(r"(sentence|phrase|word|term|claim)", text):
        return True
    return False


def _verify_rewrite(
    provider: Provider,
    model: str,
    original: str,
    revised: str,
    critique: str,
    timeout_seconds: int,
) -> tuple[bool, dict[str, Any]]:
    system_prompt = (
        "You are a careful scientific editor. Return ONLY JSON with keys: ok, fluency_score, faithfulness_score, "
        "alignment_score, issues. Scores are 0-1. ok=false if meaning changes, invented facts, or worse clarity."
    )
    user_prompt = (
        "Evaluate the rewrite for fluency, faithfulness to meaning, and alignment with the critique.\n\n"
        f"CRITIQUE:\n{critique}\n\n"
        f"ORIGINAL:\n{original}\n\n"
        f"REVISED:\n{revised}\n\n"
        "Return JSON only."
    )
    resp = provider.chat(
        ChatRequest(
            model=model,
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            temperature=0.0,
            max_tokens=320,
            timeout_seconds=timeout_seconds,
            metadata={"purpose": "rewrite_verifier"},
        )
    )
    candidate = extract_json_candidate(resp.content) or resp.content
    parsed = json.loads(candidate)
    ok = bool(parsed.get("ok", False))
    return ok, parsed


def _verdict_passes(verdict: dict[str, Any] | None) -> tuple[bool, str | None]:
    if not isinstance(verdict, dict):
        return False, "missing_verdict"
    issues = verdict.get("issues", [])
    issues_text = " ".join(str(x) for x in issues) if isinstance(issues, list) else str(issues)
    low = issues_text.lower()
    reject_markers = (
        "awkward",
        "mechanical",
        "invented",
        "meaning",
        "unsupported",
        "assumption",
        "assumptions",
        "not present in the original",
        "not explicitly stated",
        "does not address",
        "truncated",
        "incomplete",
        "verbose",
        "unnatural",
    )
    if any(marker in low for marker in reject_markers):
        return False, "verifier_issue_flag"
    try:
        if float(verdict.get("fluency_score", 1.0)) < 0.72:
            return False, "low_fluency"
        if float(verdict.get("faithfulness_score", 1.0)) < 0.72:
            return False, "low_faithfulness"
        if float(verdict.get("alignment_score", 1.0)) < 0.68:
            return False, "low_alignment"
    except Exception:
        return False, "invalid_verdict_scores"
    return bool(verdict.get("ok", False)), None if verdict.get("ok", False) else "verdict_not_ok"


def _prefer_rewrite_failure_reason(paragraph_reason: str | None, span_reason: str | None) -> str | None:
    priority = {
        "unsupported_addition": 100,
        "editorial_meta_text": 95,
        "numeric_loss": 90,
        "meaning_drift": 85,
        "punctuation_incomplete": 80,
        "incomplete_phrase": 78,
        "unbalanced_parentheses": 76,
        "unbalanced_brackets": 75,
        "markdown_heading": 74,
        "markdown_section_label": 73,
        "low_faithfulness": 72,
        "low_alignment": 71,
        "low_fluency": 70,
        "scope_drift": 68,
        "low_span_overlap": 65,
        "length_drift": 60,
        "low_overlap": 55,
        "no_change": 50,
        "empty_revision": 45,
    }
    candidates = [reason for reason in [paragraph_reason, span_reason] if reason]
    if not candidates:
        return None
    return max(candidates, key=lambda reason: priority.get(reason, 10))


def _deterministic_rewrite_for_group(group: list[dict[str, Any]], target_span: str, section: str) -> tuple[str | None, str | None]:
    issue_types = {str(c.get("issue_type", "")).strip().lower() for c in group if str(c.get("issue_type", "")).strip()}
    target = target_span.strip()
    if not target:
        return None, None
    if len(_split_sentences(target)) != 1:
        return None, None
    if len(target.split()) > 30:
        return None, None
    if target.count(",") > 2:
        return None, None
    if target.count("(") != target.count(")"):
        return None, None
    if target.count("[") != target.count("]"):
        return None, None
    if target.endswith(("Fig.", "Figure", "Table")):
        return None, None
    critique_text = " ".join(str(c.get("critique", "")) for c in group).lower()
    if issue_types <= {"evidence/overclaim concern"}:
        revised = _calibrate_claim_sentence(target)
        if revised.strip() != target:
            return revised, "claim_calibration_fallback"
        return None, None
    if issue_types <= {"clarity", "section_issue", "structure/organization"} and section in {"results", "discussion", "conclusions"}:
        revised = _calibrate_claim_sentence(target)
        if revised.strip() != target:
            return revised, "claim_calibration_fallback"
        if " broad" in f" {target.lower()} " and any(marker in critique_text for marker in ["tested evidence", "tested scope", "narrow the scope", "narrow the claim"]):
            revised = re.sub(r"\bis broad\b", "is limited to the tested evidence", target, flags=re.IGNORECASE)
            if revised.strip() != target:
                return revised, "scope_limiting_fallback"
    if issue_types <= {"clarity", "grammar/style"}:
        revised = _rewrite_candidate(target)
        if revised.strip() != target:
            return revised, "local_clarity_fallback"
        return None, None
    if issue_types <= {"section_issue", "structure/organization"} and section in {"introduction", "results", "discussion", "conclusions"}:
        revised = _rewrite_candidate(target)
        if revised.strip() != target:
            return revised, "localized_structure_fallback"
    return None, None


def _generate_suggested_changes(
    base_docx: Path,
    comments: list[dict[str, Any]],
    source_mode: dict[str, Any],
    project_id: str | None,
    run_id: str | None,
    provider: Provider | None,
    model: str | None,
    rewrite_model: str | None,
    timeout_seconds: int,
) -> tuple[list[dict[str, Any]], list[dict[str, Any]]]:
    raw_docx = Document(str(base_docx))
    raw_paragraphs = [p.text or "" for p in raw_docx.paragraphs]
    paragraphs = [normalize_review_artifact_text(text).strip() for text in raw_paragraphs]
    section_by_idx = _section_lookup_for_docx(base_docx)
    grouped: dict[int, list[dict[str, Any]]] = {}
    for c in comments:
        pidx = c.get("paragraph_index")
        if isinstance(pidx, int):
            grouped.setdefault(pidx, []).append(c)
    changes: list[dict[str, Any]] = []
    applied: list[dict[str, Any]] = []
    blocked_sections = {"front_matter", "references", "header_footer"}
    abstract_allowed = {"evidence/overclaim concern", "clarity", "structure/organization"}
    global_issue_types = {"framing", "global"}
    for pidx, paragraph_group in sorted(grouped.items(), key=lambda x: x[0]):
        section = section_by_idx.get(pidx, "body")
        original = paragraphs[pidx].strip()
        for target_span, group in _group_comments_by_rewrite_target(original, paragraph_group):
            change_id = f"chg_{uuid4().hex[:10]}"
            issue_types = sorted({str(c.get("issue_type", "")).strip() for c in group if str(c.get("issue_type", "")).strip()})
            severity = max((c.get("severity") for c in group), key=_severity_rank, default="medium")
            comment_ids = [c.get("comment_id") for c in group if c.get("comment_id")]
            base_entry = {
                "change_id": change_id,
                "project_id": project_id,
                "run_id": run_id,
                "source_mode": source_mode.get("mode"),
                "target_section": section,
                "target_paragraph_index": pidx,
                "target_span": target_span[:1200],
                "original_text": original[:1200],
                "revised_text": None,
                "originating_comment_ids": comment_ids,
                "issue_types": issue_types,
                "severity": severity,
                "rationale": None,
                "confidence": None,
                "status": "skipped",
                "skip_reason": None,
                "comment_count_in_group": len(group),
                "rewrite_strategy": "model_span_rewrite",
                "verification": None,
            }
            if issue_types and all(it.lower() in global_issue_types for it in issue_types):
                base_entry["skip_reason"] = "global_issue_not_localized"
                changes.append(base_entry)
                continue
            if _is_global_issue_not_localized(group, original):
                base_entry["skip_reason"] = "global_issue_not_localized"
                changes.append(base_entry)
                continue
            if _requires_nonlocal_addition(group, original, section):
                base_entry["skip_reason"] = "no_safe_local_rewrite"
                changes.append(base_entry)
                continue
            low = original.lower().strip()
            if section in blocked_sections or _is_front_or_back_matter_text(low):
                base_entry["skip_reason"] = "blocked_section"
                changes.append(base_entry)
                continue
            if _is_heading_paragraph(original):
                base_entry["skip_reason"] = "heading_paragraph"
                changes.append(base_entry)
                continue
            if re.match(r"^[^a-z0-9]*fig\\.|^[^a-z0-9]*figure\\b|^[^a-z0-9]*table\\b", low):
                base_entry["skip_reason"] = "caption_blocked"
                changes.append(base_entry)
                continue
            if _is_placeholder_paragraph(low):
                base_entry["skip_reason"] = "placeholder_text"
                changes.append(base_entry)
                continue
            if not original:
                base_entry["skip_reason"] = "too_short"
                changes.append(base_entry)
                continue
            if len(original.split()) < 6:
                fallback_span, fallback_strategy = _deterministic_rewrite_for_group(group, target_span, section)
                if fallback_span:
                    fallback = _apply_span_rewrite(original, target_span, fallback_span)
                    ok_fallback_span, fallback_span_reason = _basic_span_rewrite_checks(target_span, fallback_span)
                    ok_fallback, fallback_reason = _basic_rewrite_checks(original, fallback)
                    if ok_fallback_span and ok_fallback:
                        base_entry["revised_text"] = fallback[:1200]
                        base_entry["rationale"] = "Deterministic span-faithful fallback applied for a short local sentence."
                        base_entry["confidence"] = 0.4
                        base_entry["rewrite_strategy"] = fallback_strategy or "deterministic_fallback"
                        base_entry["verification"] = {"ok": True, "fallback": True, "reason": "short_sentence_safe_fallback"}
                        base_entry["status"] = "applied"
                        base_entry["skip_reason"] = None
                        changes.append(base_entry)
                        applied.append({"paragraph_index": pidx, "revised_text": fallback, "target_span": target_span})
                        continue
                    base_entry["skip_reason"] = fallback_span_reason or fallback_reason or "too_short"
                    changes.append(base_entry)
                    continue
                base_entry["skip_reason"] = "too_short"
                changes.append(base_entry)
                continue
            if section == "abstract":
                if not any(str(c.get("issue_type", "")).lower() in abstract_allowed for c in group) and not any(
                    c.get("allow_abstract", False) for c in group
                ):
                    base_entry["skip_reason"] = "abstract_high_threshold"
                    changes.append(base_entry)
                    continue
            if provider is None or (rewrite_model or model) is None:
                fallback_span, fallback_strategy = _deterministic_rewrite_for_group(group, target_span, section)
                if fallback_span:
                    fallback = _apply_span_rewrite(original, target_span, fallback_span)
                    ok_fallback_span, fallback_span_reason = _basic_span_rewrite_checks(target_span, fallback_span)
                    ok_fallback, fallback_reason = _basic_rewrite_checks(original, fallback)
                    if ok_fallback_span and ok_fallback:
                        base_entry["revised_text"] = fallback[:1200]
                        base_entry["rationale"] = "Deterministic span-faithful fallback applied because no rewrite model was available."
                        base_entry["confidence"] = 0.4
                        base_entry["rewrite_strategy"] = fallback_strategy or "deterministic_fallback"
                        base_entry["verification"] = {"ok": True, "fallback": True, "reason": "no_provider"}
                        base_entry["status"] = "applied"
                        base_entry["skip_reason"] = None
                        changes.append(base_entry)
                        applied.append({"paragraph_index": pidx, "revised_text": fallback, "target_span": target_span})
                        continue
                    base_entry["skip_reason"] = fallback_span_reason or fallback_reason or "no_provider"
                    changes.append(base_entry)
                    continue
                base_entry["skip_reason"] = "no_provider"
                changes.append(base_entry)
                continue
            prev_text = paragraphs[pidx - 1].strip() if pidx - 1 >= 0 else ""
            next_text = paragraphs[pidx + 1].strip() if pidx + 1 < len(paragraphs) else ""
            context = "\n\n".join(
                [
                    f"PREV: {prev_text}" if prev_text else "",
                    f"TARGET: {original}",
                    f"NEXT: {next_text}" if next_text else "",
                ]
            ).strip()
            critique_lines = []
            for c in group:
                critique_lines.append(
                    f"- {c.get('issue_type','')} ({c.get('severity','')}): {c.get('critique','')}. "
                    f"Suggestion: {c.get('suggested_revision','')}"
                )
            section_rules = {
                "abstract": "Keep the abstract concise; only adjust claim calibration, ambiguity, or scope alignment. Avoid adding procedural detail unless required.",
                "introduction": "Improve framing, novelty clarity, and problem statement precision without adding new claims.",
                "methods": "Improve reproducibility phrasing, parameter clarity, and disclosure without altering actual procedures.",
                "experimental": "Improve reproducibility phrasing, parameter clarity, and disclosure without altering actual procedures.",
                "results": "Clarify claim support, yield/metric precision, and interpretation boundaries; avoid overstating.",
                "discussion": "Tone down overreach and improve interpretation discipline; keep meaning intact.",
                "conclusion": "Align with demonstrated scope and avoid broad overclaims.",
                "conclusions": "Align with demonstrated scope and avoid broad overclaims.",
            }
            rule = section_rules.get(section, "Improve clarity and flow while preserving scientific meaning.")
            system_prompt = (
                "You are an expert scientific editor. Return ONLY JSON with keys: revised_text, rationale, confidence. "
                "Do not add new facts or citations. Preserve meaning unless a comment requests calibration. "
                "Make minimal, high-impact edits rather than rewriting the whole paragraph."
            )
            user_prompt = (
                f"SECTION: {section}\n"
                f"RULE: {rule}\n\n"
                f"COMMENTS:\n" + "\n".join(critique_lines) + "\n\n"
                f"CONTEXT:\n{context}\n\n"
                f'TARGET SPAN:\n{target_span}\n\n'
                "Rewrite ONLY the TARGET SPAN so it resolves the comment while preserving the paragraph's meaning. "
                "Do not rewrite surrounding sentences. Return JSON only."
            )
            use_model = rewrite_model or model
            try:
                resp = provider.chat(
                    ChatRequest(
                        model=use_model,
                        system_prompt=system_prompt,
                        user_prompt=user_prompt,
                        temperature=0.2,
                        max_tokens=700,
                        timeout_seconds=timeout_seconds,
                        metadata={
                            "purpose": "suggested_changes",
                            "project_id": project_id,
                            "run_id": run_id,
                            "section": section,
                        },
                    )
                )
                candidate = extract_json_candidate(resp.content) or resp.content
                parsed = json.loads(candidate)
            except Exception:
                base_entry["skip_reason"] = "generation_failed"
                changes.append(base_entry)
                continue
            revised_span = str(parsed.get("revised_text", "") or "").strip()
            rationale = str(parsed.get("rationale", "") or "").strip()
            confidence = parsed.get("confidence", None)
            ok_span, span_reason = _basic_span_rewrite_checks(target_span, revised_span)
            candidate_text = _apply_span_rewrite(original, target_span, revised_span)
            ok_basic, reason = _basic_rewrite_checks(original, candidate_text)
            if not ok_span or not ok_basic:
                reason = _prefer_rewrite_failure_reason(reason, span_reason)
                fallback_span, fallback_strategy = _deterministic_rewrite_for_group(group, target_span, section)
                if fallback_span:
                    fallback = _apply_span_rewrite(original, target_span, fallback_span)
                    ok_fallback_span, fallback_span_reason = _basic_span_rewrite_checks(target_span, fallback_span)
                    ok_fallback, fallback_reason = _basic_rewrite_checks(original, fallback)
                    if ok_fallback_span and ok_fallback:
                        base_entry["revised_text"] = fallback[:1200]
                        base_entry["rationale"] = "Model rewrite failed local quality checks; applied a deterministic span-faithful fallback."
                        base_entry["confidence"] = 0.45
                        base_entry["rewrite_strategy"] = fallback_strategy or "deterministic_fallback"
                        base_entry["verification"] = {"ok": True, "fallback": True, "reason": reason or "model_rewrite_rejected"}
                        base_entry["status"] = "applied"
                        base_entry["skip_reason"] = None
                        changes.append(base_entry)
                        applied.append({"paragraph_index": pidx, "revised_text": fallback, "target_span": target_span})
                        continue
                    base_entry["skip_reason"] = fallback_span_reason or fallback_reason or reason or "rewrite_rejected_low_quality"
                    changes.append(base_entry)
                    continue
                base_entry["skip_reason"] = reason or "rewrite_rejected_low_quality"
                changes.append(base_entry)
                continue
            try:
                ok_verify, verdict = _verify_rewrite(
                    provider=provider,
                    model=use_model,
                    original=original,
                    revised=candidate_text,
                    critique="; ".join(critique_lines)[:800],
                    timeout_seconds=timeout_seconds,
                )
            except Exception:
                ok_verify, verdict = False, {"ok": False, "issues": ["verifier_failed"]}
            base_entry["verification"] = verdict
            verdict_ok, verdict_reason = _verdict_passes(verdict)
            ok_verify = ok_verify and verdict_ok
            if not ok_verify:
                issues = verdict.get("issues", []) if isinstance(verdict, dict) else []
                if issues:
                    repair_system = (
                        "You are a careful scientific editor revising a draft. Fix the issues listed. "
                        "Return ONLY JSON with keys: revised_text, rationale, confidence."
                    )
                    repair_user = (
                        f"SECTION: {section}\n"
                        f"RULE: {rule}\n\n"
                        f"ISSUES TO FIX:\n{issues}\n\n"
                        f"CRITIQUE:\n" + "\n".join(critique_lines) + "\n\n"
                        f"ORIGINAL:\n{original}\n\n"
                        f"CURRENT REVISION:\n{revised_span}\n\n"
                        "Provide a corrected revision that resolves the issues. Return JSON only."
                    )
                    try:
                        repair_resp = provider.chat(
                            ChatRequest(
                                model=use_model,
                                system_prompt=repair_system,
                                user_prompt=repair_user,
                                temperature=0.2,
                                max_tokens=700,
                                timeout_seconds=timeout_seconds,
                                metadata={"purpose": "suggested_changes_revise", "section": section},
                            )
                        )
                        repair_candidate = extract_json_candidate(repair_resp.content) or repair_resp.content
                        repair_parsed = json.loads(repair_candidate)
                        revised_span = str(repair_parsed.get("revised_text", "") or "").strip()
                        rationale = str(repair_parsed.get("rationale", "") or "").strip()
                        confidence = repair_parsed.get("confidence", None)
                        ok_span, span_reason = _basic_span_rewrite_checks(target_span, revised_span)
                        candidate_text = _apply_span_rewrite(original, target_span, revised_span)
                        ok_basic, reason = _basic_rewrite_checks(original, candidate_text)
                        if ok_span and ok_basic:
                            ok_verify, verdict = _verify_rewrite(
                                provider=provider,
                                model=use_model,
                                original=original,
                                revised=candidate_text,
                                critique="; ".join(critique_lines)[:800],
                                timeout_seconds=timeout_seconds,
                            )
                            base_entry["verification"] = verdict
                            verdict_ok, verdict_reason = _verdict_passes(verdict)
                            ok_verify = ok_verify and verdict_ok
                        else:
                            ok_verify = False
                            verdict_reason = _prefer_rewrite_failure_reason(reason, span_reason)
                    except Exception:
                        ok_verify = False
                if not ok_verify:
                    fallback_span, fallback_strategy = _deterministic_rewrite_for_group(group, target_span, section)
                    if fallback_span:
                        fallback = _apply_span_rewrite(original, target_span, fallback_span)
                        ok_fallback_span, fallback_span_reason = _basic_span_rewrite_checks(target_span, fallback_span)
                        ok_fallback, fallback_reason = _basic_rewrite_checks(original, fallback)
                        if ok_fallback_span and ok_fallback:
                            rationale = "Deterministic span-faithful fallback applied after model rewrite failed quality checks."
                            confidence = 0.45
                            base_entry["verification"] = {"ok": True, "fallback": True, "reason": verdict_reason or "model_rewrite_rejected"}
                            base_entry["rewrite_strategy"] = fallback_strategy or "deterministic_fallback"
                            base_entry["revised_text"] = fallback[:1200]
                            base_entry["rationale"] = rationale
                            base_entry["confidence"] = confidence
                            base_entry["status"] = "applied"
                            base_entry["skip_reason"] = None
                            changes.append(base_entry)
                            applied.append({"paragraph_index": pidx, "revised_text": fallback, "target_span": target_span})
                            continue
                        base_entry["skip_reason"] = fallback_span_reason or fallback_reason or verdict_reason or "rewrite_rejected_low_quality"
                        changes.append(base_entry)
                        continue
                    base_entry["skip_reason"] = verdict_reason or "rewrite_rejected_low_quality"
                    changes.append(base_entry)
                    continue
            final_text = candidate_text
            base_entry["revised_text"] = final_text[:1200]
            base_entry["rationale"] = rationale[:600] if rationale else None
            try:
                base_entry["confidence"] = float(confidence) if confidence is not None else None
            except Exception:
                base_entry["confidence"] = None
            base_entry["status"] = "applied"
            base_entry["skip_reason"] = None
            changes.append(base_entry)
            applied.append({"paragraph_index": pidx, "revised_text": final_text, "target_span": target_span})
    return changes, applied


def _final_suggested_change_arbitration(
    changes: list[dict[str, Any]],
    base_docx: Path,
    provider: Provider | None,
    model: str | None,
    timeout_seconds: int,
) -> list[dict[str, Any]]:
    if not changes or provider is None or not model:
        return changes
    paragraphs = _analysis_paragraphs(base_docx)
    section_by_idx = _build_section_index_map(paragraphs)
    paper_snapshot = _paper_context_snapshot(base_docx)
    items: list[dict[str, Any]] = []
    applicable_ids: set[str] = set()
    for change in changes:
        if str(change.get("status", "")) != "applied":
            continue
        pidx = change.get("target_paragraph_index")
        if not isinstance(pidx, int) or not (0 <= pidx < len(paragraphs)):
            continue
        original = str(change.get("original_text", "") or paragraphs[pidx]).strip()
        revised = str(change.get("revised_text", "") or "").strip()
        section = section_by_idx.get(pidx, "body")
        change_id = str(change.get("change_id") or f"chg_{uuid4().hex[:10]}")
        change["change_id"] = change_id
        applicable_ids.add(change_id)
        items.append(
            {
                "change_id": change_id,
                "section": section,
                "target_span": str(change.get("target_span", ""))[:320],
                "paragraph": paragraphs[pidx][:420],
                "previous_paragraph": paragraphs[pidx - 1][:180] if pidx > 0 else "",
                "next_paragraph": paragraphs[pidx + 1][:180] if pidx + 1 < len(paragraphs) else "",
                "candidate_change": {
                    "issue_types": change.get("issue_types", []),
                    "original_text": original[:420],
                    "revised_text": revised[:420],
                    "rationale": str(change.get("rationale", "") or "")[:220],
                },
            }
        )
    if not items:
        return changes
    payload = {"paper_snapshot": paper_snapshot[:700], "items": items}
    system_prompt = (
        "You are the final editorial arbiter for manuscript suggested revisions. "
        "Judge each candidate revision in context. Return ONLY JSON with key decisions, where decisions is a list of objects with keys: "
        "change_id, action, revised_text, rationale, deletion_reason. "
        "action must be keep, revise, or drop. Drop revisions that are awkward, generic, semantically drifted, or not responsive."
    )
    try:
        resp = provider.chat(
            ChatRequest(
                model=model,
                system_prompt=system_prompt,
                user_prompt=json.dumps(payload, ensure_ascii=False),
                temperature=0.0,
                max_tokens=1400,
                timeout_seconds=timeout_seconds,
                metadata={"purpose": "final_suggested_change_arbitration_batch"},
            )
        )
        parsed = json.loads(extract_json_candidate(resp.content) or resp.content)
    except Exception:
        parsed = {"decisions": []}
    decisions = parsed.get("decisions", []) if isinstance(parsed, dict) else []
    if not decisions and len(items) == 1 and isinstance(parsed, dict):
        single = dict(parsed)
        single.setdefault("change_id", items[0]["change_id"])
        decisions = [single]
    decisions_by_id = {
        str(item.get("change_id")): item
        for item in decisions
        if isinstance(item, dict) and str(item.get("change_id", "")).strip()
    }
    revised_changes: list[dict[str, Any]] = []
    for change in changes:
        if str(change.get("change_id", "")) not in applicable_ids:
            revised_changes.append(change)
            continue
        original = str(change.get("original_text", "")).strip()
        parsed = decisions_by_id.get(str(change.get("change_id", "")), {"action": "keep"})
        action = str(parsed.get("action", "keep")).strip().lower()
        if action == "drop":
            dropped = dict(change)
            dropped["status"] = "skipped"
            dropped["skip_reason"] = "final_audit_rejected"
            revised_changes.append(dropped)
            continue
        candidate = dict(change)
        if action == "revise":
            revised_text = str(parsed.get("revised_text", "") or "").strip()
            rationale = str(parsed.get("rationale", "") or "").strip()
            ok_basic, reason = _basic_rewrite_checks(original, revised_text)
            if ok_basic:
                if provider is not None and model:
                    ok_verify, verdict = _verify_rewrite(
                        provider=provider,
                        model=model,
                        original=original,
                        revised=revised_text,
                        critique="; ".join(str(x) for x in candidate.get("issue_types", []))[:800],
                        timeout_seconds=timeout_seconds,
                    )
                    verdict_ok, verdict_reason = _verdict_passes(verdict)
                    if ok_verify and verdict_ok:
                        candidate["revised_text"] = revised_text[:1200]
                        candidate["rationale"] = rationale or candidate.get("rationale")
                        candidate["verification"] = verdict
                    else:
                        candidate["status"] = "skipped"
                        candidate["skip_reason"] = verdict_reason or "final_audit_rejected"
                else:
                    candidate["revised_text"] = revised_text[:1200]
                    candidate["rationale"] = rationale or candidate.get("rationale")
            else:
                candidate["status"] = "skipped"
                candidate["skip_reason"] = reason or "final_audit_rejected"
        revised_changes.append(candidate)
    return revised_changes


def build_annotated_manuscript_output(
    source_path: Path,
    doc: ParsedDocument,
    review: Any,
    output_dir: Path,
    project_id: str | None = None,
    run_id: str | None = None,
    provider: Provider | None = None,
    model: str | None = None,
    rewrite_model: str | None = None,
    comment_audit_model: str | None = None,
    suggestion_audit_model: str | None = None,
    timeout_seconds: int = 120,
) -> dict[str, Any]:
    source_mode = detect_source_mode(source_path)
    if source_mode["mode"] == "original_docx":
        base_docx = source_path
        reviewed_name = "reviewed_manuscript_with_comments.docx"
        suggested_name = "reviewed_manuscript_with_suggested_changes.docx"
    else:
        base_docx = output_dir / "surrogate_manuscript_from_pdf_base.docx"
        create_docx_from_plain_text(doc.cleaned_text, base_docx, title=source_path.stem)
        reviewed_name = "surrogate_manuscript_from_pdf_with_comments.docx"
        suggested_name = "surrogate_manuscript_from_pdf_with_suggested_changes.docx"
    comments = review_to_comment_entries(review, doc=doc, base_docx=base_docx)
    comments = _assign_paragraph_indices(comments, base_docx)
    comments = _localize_comment_entries(comments, base_docx)
    comments = _revise_comment_entries(comments, base_docx)
    comments = _enrich_comment_suggestions(comments, base_docx)
    comments = _dedupe_comment_entries(comments)
    comments = _filter_comment_entries_by_paragraph_quality(comments, base_docx)
    comments = _balance_comment_entries(comments, base_docx, max_comments=36)
    comments = _limit_comments_per_paragraph(comments, max_per_paragraph=2)
    comments = _final_comment_arbitration(
        comments,
        base_docx=base_docx,
        provider=provider,
        model=comment_audit_model or rewrite_model or model,
        timeout_seconds=timeout_seconds,
    )
    comments = _dedupe_comment_entries(comments)
    comments = _filter_comment_entries_by_paragraph_quality(comments, base_docx)
    comments = _balance_comment_entries(comments, base_docx, max_comments=36)
    comments = _limit_comments_per_paragraph(comments, max_per_paragraph=2)
    if not comments:
        fallback_idx = _first_substantive_paragraph_index(base_docx)
        comments = [
            {
                "comment_id": f"cmt_{uuid4().hex[:10]}",
                "paragraph_index": fallback_idx,
                "issue_type": "clarity",
                "severity": "medium",
                "critique": "The first substantive sentence still needs a clearer editorial focus. State one concrete claim and one specific supporting detail instead of a broad setup sentence.",
                "suggested_revision": "Suggested wording direction: keep the main claim in the opening sentence and move any qualification or background phrase into the next sentence.",
                "rationale": "Fallback after quality gates removed low-value comment candidates.",
            }
        ]
    reviewed_docx = output_dir / reviewed_name
    result = create_commented_docx_copy(base_docx, reviewed_docx, comments, comment_tag=run_id)
    validation = validate_commented_docx(base_docx, reviewed_docx)
    section_map = _section_map_for_docx(base_docx)
    changes_manifest, applied_changes = _generate_suggested_changes(
        base_docx=base_docx,
        comments=comments,
        source_mode=source_mode,
        project_id=project_id,
        run_id=run_id,
        provider=provider,
        model=model,
        rewrite_model=rewrite_model,
        timeout_seconds=timeout_seconds,
    )
    changes_manifest = _final_suggested_change_arbitration(
        changes_manifest,
        base_docx=base_docx,
        provider=provider,
        model=suggestion_audit_model or comment_audit_model or rewrite_model or model,
        timeout_seconds=timeout_seconds,
    )
    suggested_docx = output_dir / suggested_name
    suggested_result = create_suggested_changes_docx(
        source_docx=base_docx,
        output_docx=suggested_docx,
        changes=[
            {
                "paragraph_index": entry.get("target_paragraph_index"),
                "revised_text": entry.get("revised_text"),
                "original_text": entry.get("original_text"),
                "status": entry.get("status"),
            }
            for entry in changes_manifest
        ],
    )
    suggested_validation = validate_suggested_changes_docx(base_docx, suggested_docx)
    base_doc = Document(str(base_docx))
    suggested_doc = Document(str(suggested_docx))
    section_by_idx = _section_lookup_for_docx(base_docx)
    front_matter_changed = 0
    references_changed = 0
    for idx, base_p in enumerate(base_doc.paragraphs):
        sec = section_by_idx.get(idx, "body")
        if idx >= len(suggested_doc.paragraphs):
            continue
        if base_p.text != suggested_doc.paragraphs[idx].text:
            if sec == "front_matter":
                front_matter_changed += 1
            if sec == "references":
                references_changed += 1
    manifest_path = output_dir / "manuscript_suggested_changes_manifest.json"
    manifest_path.write_text(json.dumps(changes_manifest, indent=2))
    validation_path = output_dir / "suggested_changes_validation.json"
    validation_payload = {
        "suggested_docx": str(suggested_docx),
        "manifest_path": str(manifest_path),
        "changes_proposed": len(changes_manifest),
        "changes_applied": suggested_result.get("changes_applied", 0),
        "follow_up_changes_applied": suggested_result.get("follow_up_changes_applied", 0),
        "applied_paragraph_indices": suggested_result.get("applied_paragraph_indices", []),
        "front_matter_changed_count": front_matter_changed,
        "references_changed_count": references_changed,
        "front_matter_untouched": front_matter_changed == 0,
        "references_untouched": references_changed == 0,
        **suggested_validation,
    }
    validation_path.write_text(json.dumps(validation_payload, indent=2))
    preannotated_policy = {
        "preserve_existing_comments": True,
        "preserve_visible_suggested_blocks": True,
        "strip_visible_suggested_blocks_from_analysis_text": True,
        "layer_new_comments_on_top": True,
        "layer_new_suggested_blocks_as_follow_up_when_needed": True,
        "dedupe_against_existing_comments": False,
        "summarize_existing_comments_as_context": False,
    }
    source_mode_artifact = {
        "project_id": project_id,
        "manuscript_source_path": str(source_path),
        "source_mode": source_mode["mode"],
        "annotation_state": source_mode.get("annotation_state"),
        "docx_annotation_state": source_mode.get("docx_annotation_state"),
        "surrogate_base_path": str(base_docx) if source_mode["mode"] != "original_docx" else None,
        "preannotated_docx_policy": preannotated_policy if source_mode["mode"] == "original_docx" else None,
        "annotation_policy": (
            "Preserve existing DOCX comments and visible suggested-change blocks; "
            "strip visible AI-Reviewer suggestion blocks from analysis text; layer new comments and follow-up suggested changes on top."
        ),
        "notes": [],
    }
    return {
        "source_mode": source_mode,
        "source_mode_artifact": source_mode_artifact,
        "base_docx": str(base_docx),
        "reviewed_docx": str(reviewed_docx),
        "comments_requested": len(comments),
        "comments_added": result.get("comments_added", 0),
        "existing_comments_before": result.get("existing_comments_before", 0),
        "existing_comments_after": result.get("existing_comments_after", 0),
        "new_ai_reviewer_comments_added_count": validation.get("new_ai_reviewer_comments_added_count", 0),
        "anchored_paragraph_indices": result.get("anchored_paragraph_indices", []),
        "comment_targets": comments,
        "validation": validation,
        "preannotated_docx_policy": preannotated_policy if source_mode["mode"] == "original_docx" else None,
        "suggested_changes_docx": str(suggested_docx),
        "suggested_changes_manifest": str(manifest_path),
        "suggested_changes_validation": validation_payload,
        "section_map": section_map,
    }
