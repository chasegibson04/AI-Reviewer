from __future__ import annotations

import difflib
import re
import shutil
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document

try:
    from docx2python import docx2python
except Exception:  # pragma: no cover
    docx2python = None

try:
    import mammoth
except Exception:  # pragma: no cover
    mammoth = None


W_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
XML_NS = "http://www.w3.org/XML/1998/namespace"
SUGGESTED_CHANGE_MARKER = "[Suggested change]"
FOLLOWUP_SUGGESTED_CHANGE_MARKER = "[Suggested change - follow-up]"

ET.register_namespace("w", W_NS["w"])


def normalize_review_artifact_text(text: str) -> str:
    if not text:
        return ""
    lines: list[str] = []
    for line in text.splitlines():
        stripped = line.strip()
        if stripped.startswith(SUGGESTED_CHANGE_MARKER):
            continue
        if stripped.startswith(FOLLOWUP_SUGGESTED_CHANGE_MARKER):
            continue
        lines.append(line)
    cleaned = "\n".join(lines)
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def get_docx_paragraph_texts(path: Path, normalize_review_artifacts: bool = False) -> list[str]:
    doc = Document(str(path))
    texts = [p.text or "" for p in doc.paragraphs]
    if normalize_review_artifacts:
        return [normalize_review_artifact_text(text) for text in texts]
    return texts


def inspect_docx_annotation_state(path: Path) -> dict:
    doc = Document(str(path))
    paragraphs = [p.text or "" for p in doc.paragraphs]
    normalized = [normalize_review_artifact_text(text) for text in paragraphs]
    visible_suggested_blocks = sum(text.count(SUGGESTED_CHANGE_MARKER) for text in paragraphs) + sum(
        text.count(FOLLOWUP_SUGGESTED_CHANGE_MARKER) for text in paragraphs
    )
    comment_texts: list[str] = []
    comment_authors: list[str] = []
    comments = list(doc.comments) if hasattr(doc, "comments") else []
    for comment in comments:
        try:
            comment_texts.append((comment.text or "").strip())
            comment_authors.append((comment.author or "").strip())
        except Exception:
            continue
    ai_reviewer_comment_count = sum(1 for author in comment_authors if "ai-reviewer" in author.lower())
    ai_reviewer_comment_count += sum(1 for text in comment_texts if "issue:" in text.lower() and "suggested revision:" in text.lower())
    non_ai_comment_count = max(0, len(comment_texts) - ai_reviewer_comment_count) if comment_texts else max(0, len(comments) - ai_reviewer_comment_count)

    track_changes = {"insertions": 0, "deletions": 0, "moves": 0}
    comments_xml_present = False
    try:
        with zipfile.ZipFile(str(path), "r") as zf:
            document_xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
            track_changes["insertions"] = len(re.findall(r"<w:ins\b", document_xml))
            track_changes["deletions"] = len(re.findall(r"<w:del\b", document_xml))
            track_changes["moves"] = len(re.findall(r"<w:move(To|From)\b", document_xml))
            comments_xml_present = "word/comments.xml" in zf.namelist()
            if comments_xml_present:
                root = ET.fromstring(zf.read("word/comments.xml"))
                extracted_authors: list[str] = []
                extracted_texts: list[str] = []
                for comment in root.findall(".//w:comment", W_NS):
                    author = comment.attrib.get(f"{{{W_NS['w']}}}author", "")
                    if author:
                        extracted_authors.append(author)
                    parts = [node.text or "" for node in comment.findall(".//w:t", W_NS)]
                    joined = "".join(parts).strip()
                    if joined:
                        extracted_texts.append(joined)
                if extracted_authors:
                    comment_authors = extracted_authors
                if extracted_texts:
                    comment_texts = extracted_texts
                    ai_reviewer_comment_count = sum(
                        1 for text in extracted_texts if "issue:" in text.lower() and "suggested revision:" in text.lower()
                    ) + sum(1 for author in comment_authors if "ai-reviewer" in author.lower())
    except Exception:
        pass

    track_change_count = sum(track_changes.values())
    if not comments and not visible_suggested_blocks and track_change_count == 0:
        annotation_state = "clean_native_docx"
    elif ai_reviewer_comment_count > 0 or visible_suggested_blocks > 0:
        annotation_state = "prior_ai_reviewer_annotated_docx" if comments or visible_suggested_blocks else "prior_ai_reviewer_comments_only"
    elif comments and track_change_count:
        annotation_state = "mixed_annotated_docx"
    elif comments:
        annotation_state = "docx_with_existing_comments"
    elif visible_suggested_blocks:
        annotation_state = "docx_with_existing_suggested_blocks"
    elif track_change_count:
        annotation_state = "docx_with_tracked_changes"
    else:
        annotation_state = "clean_native_docx"

    return {
        "annotation_state": annotation_state,
        "existing_comments_count": len(comment_texts) if comment_texts else len(comments),
        "existing_comment_authors": sorted({a for a in comment_authors if a})[:50],
        "existing_comment_excerpts": [text[:200] for text in comment_texts[:10]],
        "existing_ai_reviewer_comment_count": ai_reviewer_comment_count,
        "existing_non_ai_comment_count": non_ai_comment_count,
        "existing_suggested_change_blocks": visible_suggested_blocks,
        "tracked_change_elements": track_changes,
        "tracked_change_total": track_change_count,
        "comments_xml_present": comments_xml_present,
        "review_artifact_text_blocks_removed_for_analysis": sum(1 for raw, clean in zip(paragraphs, normalized) if raw.strip() != clean.strip()),
    }


def parse_docx_structured(path: Path) -> dict:
    doc = Document(str(path))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.lower().startswith("heading")]
    warnings: list[str] = []

    alt_text = ""
    if docx2python is not None:
        try:
            alt = docx2python(str(path))
            alt_text = "\n".join(alt.body_runs)
        except Exception as exc:
            warnings.append(f"docx2python_failed:{exc}")

    html_text = ""
    if mammoth is not None:
        try:
            with path.open("rb") as f:
                html_text = mammoth.convert_to_html(f).value
        except Exception as exc:
            warnings.append(f"mammoth_failed:{exc}")

    comment_count = len(list(doc.comments)) if hasattr(doc, "comments") else 0
    annotation_state = inspect_docx_annotation_state(path)
    return {
        "tool": "python-docx",
        "paragraph_count": len(paragraphs),
        "headings": headings[:300],
        "text": "\n".join(paragraphs),
        "comments_count": comment_count,
        "annotation_state": annotation_state,
        "alt_text_excerpt": alt_text[:4000],
        "html_excerpt": html_text[:4000],
        "warnings": warnings,
    }


def create_commented_docx_copy(
    source_docx: Path,
    output_docx: Path,
    comments: list[dict],
    author: str = "AI-Reviewer",
    initials: str = "AIR",
    comment_tag: str | None = None,
) -> dict:
    doc = Document(str(source_docx))
    before_state = inspect_docx_annotation_state(source_docx)
    paragraph_map: dict[int, list[dict]] = {}
    for entry in comments:
        idx = int(entry.get("paragraph_index", -1))
        if idx < 0:
            continue
        paragraph_map.setdefault(idx, []).append(entry)

    added = 0
    anchored_paragraph_indices: list[int] = []

    def _anchor_runs(paragraph, anchor_text: str):
        full_text = paragraph.text or ""
        if not anchor_text or not full_text:
            return None, None
        low_full = full_text.lower()
        low_anchor = anchor_text.lower().strip()
        pos = low_full.find(low_anchor)
        if pos < 0:
            return None, None
        before = full_text[:pos]
        middle = full_text[pos : pos + len(anchor_text)]
        after = full_text[pos + len(anchor_text) :]
        paragraph.text = ""
        before_run = paragraph.add_run(before) if before else None
        middle_run = paragraph.add_run(middle)
        after_run = paragraph.add_run(after) if after else None
        return middle_run, middle_run or before_run or after_run

    for pidx, entries in paragraph_map.items():
        if pidx >= len(doc.paragraphs):
            continue
        paragraph = doc.paragraphs[pidx]
        if not paragraph.runs:
            paragraph.add_run(" ")
        for item in entries[:3]:
            first_run = paragraph.runs[0]
            last_run = paragraph.runs[-1]
            anchor_text = str(item.get("anchor_text", "")).strip()
            anchored_first, anchored_last = _anchor_runs(paragraph, anchor_text)
            if anchored_first is not None and anchored_last is not None:
                first_run, last_run = anchored_first, anchored_last
            elif paragraph.runs:
                first_run, last_run = paragraph.runs[0], paragraph.runs[-1]
            body_parts = [
                f"Review round: {comment_tag}" if comment_tag else "",
                f"Issue: {item.get('issue_type', 'general')}",
                f"Severity: {item.get('severity', 'medium')}",
                f"Critique: {item.get('critique', '')}",
                f"Suggested revision: {item.get('suggested_revision', '')}",
                f"Rationale: {item.get('rationale', '')}",
            ]
            evidence = item.get("evidence_source")
            if evidence:
                body_parts.append(f"Evidence: {evidence}")
            quote = item.get("manuscript_quote")
            if quote:
                body_parts.append(f"Quote: {quote}")
            
            body = "\n".join([p for p in body_parts if p]).strip()
            comment = doc.comments.add_comment(text=body, author=author, initials=initials)
            first_run.mark_comment_range(last_run, comment.comment_id)
            added += 1
            anchored_paragraph_indices.append(pidx)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    after_state = inspect_docx_annotation_state(output_docx)
    return {
        "output": str(output_docx),
        "comments_added": added,
        "anchored_paragraph_indices": anchored_paragraph_indices,
        "existing_comments_before": before_state.get("existing_comments_count", 0),
        "existing_comments_after": after_state.get("existing_comments_count", 0),
        "annotation_state_before": before_state,
        "annotation_state_after": after_state,
    }


def create_docx_from_plain_text(text: str, output_docx: Path, title: str = "Manuscript") -> Path:
    doc = Document()
    doc.add_heading(title, level=1)

    def _normalize_heading(raw: str) -> str:
        cleaned = re.sub(r"^[■*\[\]]+", "", raw).strip()
        cleaned = re.sub(r"\s+", " ", cleaned)
        return cleaned.strip(":").strip()

    def _is_heading_block(block: str) -> bool:
        if len(block) > 80:
            return False
        low = block.lower().strip(":").strip()
        if low in {
            "abstract",
            "introduction",
            "experimental",
            "methods",
            "results",
            "discussion",
            "conclusion",
            "conclusions",
            "author contributions",
            "funding",
            "associated content",
            "author information",
            "abbreviations",
            "references",
            "keywords",
        }:
            return True
        if block.isupper() and len(block.split()) <= 6:
            return True
        if re.match(r"^\[?[A-Z][A-Z\s-]{3,}\]?$", block):
            return True
        return False

    for block in [b.strip() for b in text.split("\n\n") if b.strip()]:
        if _is_heading_block(block):
            doc.add_heading(_normalize_heading(block), level=2)
        else:
            doc.add_paragraph(block)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    return output_docx


def extract_docx_body_text(path: Path) -> str:
    doc = Document(str(path))
    return "\n".join(p.text for p in doc.paragraphs)


def _preserve_space(node: ET.Element, text: str) -> None:
    if text.startswith(" ") or text.endswith(" ") or "  " in text or "\n" in text or "\t" in text:
        node.set(f"{{{XML_NS}}}space", "preserve")


def _append_text_run(parent: ET.Element, text: str) -> None:
    if not text:
        return
    run = ET.SubElement(parent, f"{{{W_NS['w']}}}r")
    text_el = ET.SubElement(run, f"{{{W_NS['w']}}}t")
    _preserve_space(text_el, text)
    text_el.text = text


def _append_tracked_change(parent: ET.Element, kind: str, text: str, change_id: int, author: str, date_iso: str) -> None:
    if not text:
        return
    tag = f"{{{W_NS['w']}}}{'ins' if kind == 'ins' else 'del'}"
    wrapper = ET.SubElement(
        parent,
        tag,
        {
            f"{{{W_NS['w']}}}id": str(change_id),
            f"{{{W_NS['w']}}}author": author,
            f"{{{W_NS['w']}}}date": date_iso,
        },
    )
    run = ET.SubElement(wrapper, f"{{{W_NS['w']}}}r")
    text_tag = f"{{{W_NS['w']}}}{'t' if kind == 'ins' else 'delText'}"
    text_el = ET.SubElement(run, text_tag)
    _preserve_space(text_el, text)
    text_el.text = text


def _tokenize_for_track_changes(text: str) -> list[str]:
    return re.findall(r"\s+|[^\s]+", text or "", flags=re.MULTILINE)


def _replace_paragraph_with_track_changes(paragraph_el: ET.Element, original: str, revised: str, change_seed: int, author: str, date_iso: str) -> int:
    original_tokens = _tokenize_for_track_changes(original)
    revised_tokens = _tokenize_for_track_changes(revised)
    matcher = difflib.SequenceMatcher(a=original_tokens, b=revised_tokens)
    preserved = [child for child in list(paragraph_el) if child.tag == f"{{{W_NS['w']}}}pPr"]
    for child in list(paragraph_el):
        if child not in preserved:
            paragraph_el.remove(child)
    change_count = 0
    for opcode, i1, i2, j1, j2 in matcher.get_opcodes():
        if opcode == "equal":
            _append_text_run(paragraph_el, "".join(original_tokens[i1:i2]))
            continue
        if opcode in {"replace", "delete"}:
            deleted = "".join(original_tokens[i1:i2])
            if deleted:
                _append_tracked_change(paragraph_el, "del", deleted, change_seed + change_count, author, date_iso)
                change_count += 1
        if opcode in {"replace", "insert"}:
            inserted = "".join(revised_tokens[j1:j2])
            if inserted:
                _append_tracked_change(paragraph_el, "ins", inserted, change_seed + change_count, author, date_iso)
                change_count += 1
    if change_count == 0:
        _append_text_run(paragraph_el, revised or original)
    return change_count


def validate_commented_docx(base_docx: Path, reviewed_docx: Path) -> dict:
    base = Document(str(base_docx))
    reviewed = Document(str(reviewed_docx))
    base_text = "\n".join(p.text for p in base.paragraphs)
    reviewed_text = "\n".join(p.text for p in reviewed.paragraphs)
    comments = list(reviewed.comments) if hasattr(reviewed, "comments") else []
    before_state = inspect_docx_annotation_state(base_docx)
    after_state = inspect_docx_annotation_state(reviewed_docx)
    comment_ranges = 0
    try:
        with zipfile.ZipFile(str(reviewed_docx), "r") as zf:
            xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
            comment_ranges = len(re.findall(r"<w:commentRangeStart\b", xml))
    except Exception:
        comment_ranges = 0
    input_preannotated = before_state.get("annotation_state") != "clean_native_docx"
    new_ai_comments_added = max(
        0,
        int(after_state.get("existing_ai_reviewer_comment_count", 0)) - int(before_state.get("existing_ai_reviewer_comment_count", 0)),
    )
    return {
        "base_docx": str(base_docx),
        "reviewed_docx": str(reviewed_docx),
        "comment_count": len(comments),
        "comment_ranges_detected": comment_ranges,
        "comments_attached": comment_ranges > 0,
        "body_text_unchanged": base_text == reviewed_text,
        "input_annotation_state": before_state,
        "output_annotation_state": after_state,
        "new_comments_added_count": max(
            0,
            int(after_state.get("existing_comments_count", 0)) - int(before_state.get("existing_comments_count", 0)),
        ),
        "new_ai_reviewer_comments_added_count": new_ai_comments_added,
        "existing_comments_preserved": int(after_state.get("existing_comments_count", 0))
        >= int(before_state.get("existing_comments_count", 0)),
        "input_preannotated": input_preannotated,
        "silent_noop_suspected": input_preannotated and new_ai_comments_added == 0,
        "meaningful_new_review_state": new_ai_comments_added > 0,
    }


def create_suggested_changes_docx(
    source_docx: Path,
    output_docx: Path,
    changes: list[dict],
) -> dict:
    doc = Document(str(source_docx))
    before_state = inspect_docx_annotation_state(source_docx)
    applied = 0
    applied_indices: list[int] = []
    tracked_change_elements_added = 0
    paragraph_rewrites: dict[int, tuple[str, str]] = {}
    for change in changes:
        if change.get("status") != "applied":
            continue
        idx = int(change.get("paragraph_index", -1))
        revised = (change.get("revised_text") or "").strip()
        if idx < 0 or not revised:
            continue
        if idx >= len(doc.paragraphs):
            continue
        existing_text = normalize_review_artifact_text(doc.paragraphs[idx].text or "").strip()
        original_text = normalize_review_artifact_text(change.get("original_text") or existing_text).strip()
        if not original_text:
            original_text = existing_text
        if not original_text or revised == original_text:
            continue
        paragraph_rewrites[idx] = (original_text, revised)
        applied += 1
        applied_indices.append(idx)
    output_docx.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(source_docx, output_docx)
    if paragraph_rewrites:
        with zipfile.ZipFile(str(output_docx), "r") as zf:
            archive = {name: zf.read(name) for name in zf.namelist()}
        root = ET.fromstring(archive["word/document.xml"])
        paragraphs_xml = root.findall(".//w:body/w:p", W_NS)
        date_iso = datetime.now(timezone.utc).replace(microsecond=0).isoformat()
        change_seed = 1000
        for idx in sorted(paragraph_rewrites):
            if idx < 0 or idx >= len(paragraphs_xml):
                continue
            original_text, revised = paragraph_rewrites[idx]
            tracked_change_elements_added += _replace_paragraph_with_track_changes(
                paragraphs_xml[idx],
                original_text,
                revised,
                change_seed + (idx * 10),
                "AI-Reviewer",
                date_iso,
            )
        archive["word/document.xml"] = ET.tostring(root, encoding="utf-8", xml_declaration=True)
        with zipfile.ZipFile(str(output_docx), "w", compression=zipfile.ZIP_DEFLATED) as zf:
            for name, payload in archive.items():
                zf.writestr(name, payload)
    after_state = inspect_docx_annotation_state(output_docx)
    return {
        "output": str(output_docx),
        "changes_applied": applied,
        "applied_paragraph_indices": applied_indices,
        "follow_up_changes_applied": 0,
        "tracked_change_elements_added": tracked_change_elements_added,
        "existing_suggested_change_blocks_before": before_state.get("existing_suggested_change_blocks", 0),
        "existing_suggested_change_blocks_after": after_state.get("existing_suggested_change_blocks", 0),
    }


def validate_suggested_changes_docx(base_docx: Path, suggested_docx: Path) -> dict:
    base = Document(str(base_docx))
    suggested = Document(str(suggested_docx))
    base_paragraphs = [p.text for p in base.paragraphs]
    suggested_paragraphs = [p.text for p in suggested.paragraphs]
    before_state = inspect_docx_annotation_state(base_docx)
    after_state = inspect_docx_annotation_state(suggested_docx)
    input_preannotated = before_state.get("annotation_state") != "clean_native_docx"
    new_blocks_added = max(
        0,
        int(after_state.get("existing_suggested_change_blocks", 0)) - int(before_state.get("existing_suggested_change_blocks", 0)),
    )
    new_tracked_changes_added = max(
        0,
        int(after_state.get("tracked_change_total", 0)) - int(before_state.get("tracked_change_total", 0)),
    )
    return {
        "base_docx": str(base_docx),
        "suggested_docx": str(suggested_docx),
        "paragraph_count_base": len(base_paragraphs),
        "paragraph_count_suggested": len(suggested_paragraphs),
        "structure_intact": len(base_paragraphs) == len(suggested_paragraphs),
        "docx_opens": True,
        "input_annotation_state": before_state,
        "output_annotation_state": after_state,
        "new_suggested_change_blocks_added": new_blocks_added,
        "new_tracked_change_elements_added": new_tracked_changes_added,
        "track_changes_present": int(after_state.get("tracked_change_total", 0)) > 0,
        "input_preannotated": input_preannotated,
        "silent_noop_suspected": input_preannotated and new_blocks_added == 0 and new_tracked_changes_added == 0,
        "meaningful_new_review_state": new_blocks_added > 0 or new_tracked_changes_added > 0,
    }
