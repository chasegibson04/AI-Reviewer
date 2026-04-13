from pathlib import Path
import zipfile
import re

from docx import Document

from ai_reviewer.config import load_config
from ai_reviewer.ingest.loaders import parse_file
from ai_reviewer.review.manuscript_annotation import build_annotated_manuscript_output, detect_source_mode
from ai_reviewer.tools.availability import scan_tool_availability
from ai_reviewer.tools.docx_tools import (
    FOLLOWUP_SUGGESTED_CHANGE_MARKER,
    create_commented_docx_copy,
    create_suggested_changes_docx,
    inspect_docx_annotation_state,
    normalize_review_artifact_text,
    parse_docx_structured,
    validate_commented_docx,
    validate_suggested_changes_docx,
)


def _count_track_changes(path: Path) -> tuple[int, int]:
    with zipfile.ZipFile(str(path), "r") as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    return len(re.findall(r"<w:ins\b", xml)), len(re.findall(r"<w:del\b", xml))


def _add_existing_comment(doc: Document, paragraph_index: int, text: str, author: str = "Editor", initials: str = "ED") -> None:
    paragraph = doc.paragraphs[paragraph_index]
    if not paragraph.runs:
        paragraph.add_run(" ")
    comment = doc.comments.add_comment(text=text, author=author, initials=initials)
    paragraph.runs[0].mark_comment_range(paragraph.runs[-1], comment.comment_id)


def test_tool_availability_has_required_entries():
    avail = scan_tool_availability()
    assert "python_docx" in avail
    assert "pymupdf" in avail
    assert "pypdf" in avail


def test_docx_parse_and_comment_copy(tmp_path: Path):
    source = tmp_path / "source.docx"
    d = Document()
    d.add_heading("Title", level=1)
    d.add_paragraph("This is paragraph one.")
    d.add_paragraph("This is paragraph two.")
    d.save(str(source))

    parsed = parse_docx_structured(source)
    assert parsed["paragraph_count"] >= 2
    out = tmp_path / "commented.docx"
    result = create_commented_docx_copy(
        source,
        out,
        [
            {
                "paragraph_index": 1,
                "issue_type": "clarity",
                "severity": "high",
                "critique": "Sentence is vague",
                "suggested_revision": "Use concrete wording.",
                "rationale": "Improve readability",
            }
        ],
    )
    assert out.exists()
    assert result["comments_added"] >= 1
    reopened = Document(str(out))
    assert len(list(reopened.comments)) >= 1


def test_suggested_changes_docx_applies_edits(tmp_path: Path):
    source = tmp_path / "source.docx"
    d = Document()
    d.add_heading("Title", level=1)
    d.add_paragraph("Paragraph one.")
    d.add_paragraph("Paragraph two.")
    d.save(str(source))

    out = tmp_path / "suggested.docx"
    result = create_suggested_changes_docx(
        source_docx=source,
        output_docx=out,
        changes=[
            {
                "paragraph_index": 1,
                "revised_text": "Paragraph one updated.",
                "status": "applied",
            }
        ],
    )
    assert out.exists()
    assert result["changes_applied"] == 1
    ins_count, del_count = _count_track_changes(out)
    assert ins_count >= 1
    assert del_count >= 1
    validation = validate_suggested_changes_docx(source, out)
    assert validation["structure_intact"] is True
    assert validation["track_changes_present"] is True
    assert validation["meaningful_new_review_state"] is True


def test_scholarly_lookup_offline_no_network():
    cfg = load_config()
    assert cfg.defaults.strict_offline is True


def test_source_mode_detection():
    detected = detect_source_mode(Path("x.docx"))
    assert detected["mode"] == "original_docx"
    assert detect_source_mode(Path("x.pdf"))["mode"] == "pdf_only_surrogate"


def test_surrogate_annotation_preserves_text(tmp_path: Path):
    src = tmp_path / "sample.md"
    src.write_text("# T\n\nParagraph one.\n\nParagraph two.", encoding="utf-8")
    parsed = parse_file(src)

    class _R:
        section_specific_comments = []
        extracted_action_items = []
        detailed_reviewer_comments = ["Improve phrasing in intro."]
        methodological_concerns = []
        novelty_concerns = []
        citation_reference_concerns = []
        writing_organization_concerns = []
        figure_table_concerns = []
        major_weaknesses = []

    out = build_annotated_manuscript_output(src, parsed, _R(), tmp_path)
    assert Path(out["reviewed_docx"]).exists()
    assert Path(out["suggested_changes_docx"]).exists()
    assert Path(out["suggested_changes_manifest"]).exists()
    assert out["validation"]["comment_count"] >= 1
    assert out["validation"]["body_text_unchanged"] is True
    assert out["validation"]["comments_attached"] is True
    assert out["source_mode_artifact"]["source_mode"] == "surrogate_other_source"
    assert out["suggested_changes_validation"]["docx_opens"] is True


def test_docx_annotation_state_detects_existing_comments_and_suggested_blocks(tmp_path: Path):
    source = tmp_path / "annotated.docx"
    d = Document()
    d.add_heading("Title", level=1)
    d.add_paragraph("Abstract")
    d.add_paragraph("This paragraph introduces the manuscript claim.")
    d.add_paragraph("This paragraph already contains a visible suggestion.\n[Suggested change] Tighter wording here.")
    _add_existing_comment(d, 2, "Existing reviewer note on the introduction.")
    d.save(str(source))

    state = inspect_docx_annotation_state(source)
    assert state["existing_comments_count"] >= 1
    assert state["existing_suggested_change_blocks"] >= 1
    assert state["annotation_state"] in {"prior_ai_reviewer_annotated_docx", "mixed_annotated_docx"}
    assert state["existing_non_ai_comment_count"] >= 1

    parsed = parse_docx_structured(source)
    assert parsed["annotation_state"]["existing_comments_count"] >= 1


def test_preannotated_docx_review_adds_new_comments_and_preserves_existing(tmp_path: Path):
    source = tmp_path / "source.docx"
    d = Document()
    d.add_heading("Introduction", level=1)
    d.add_paragraph("Recently, it has become increasingly possible to predict reaction outcomes, although the available data remain limited.")
    d.add_heading("Methods", level=1)
    d.add_paragraph("The automated generation of reaction arrays to optimize a coupling reaction is a contemporary problem with multiple decision points.")
    _add_existing_comment(d, 1, "Existing editor note on framing.")
    d.save(str(source))

    parsed = parse_file(source)

    class _Item:
        def __init__(self, action: str, priority: str = "medium"):
            self.action = action
            self.priority = priority

    class _Sec:
        def __init__(self, section: str, comment: str, severity: str = "medium"):
            self.section = section
            self.comment = comment
            self.severity = severity

    class _R:
        section_specific_comments = [_Sec("Introduction", "Clarify the introductory framing claim.")]
        extracted_action_items = [_Item("Split the overloaded methods sentence and make the constraint explicit.")]
        detailed_reviewer_comments = []
        methodological_concerns = []
        novelty_concerns = []
        citation_reference_concerns = []
        writing_organization_concerns = []
        figure_table_concerns = []
        major_weaknesses = []

    out = build_annotated_manuscript_output(source, parsed, _R(), tmp_path, run_id="testround")
    assert out["source_mode"]["annotation_state"] in {"docx_with_existing_comments", "mixed_annotated_docx"}
    assert out["existing_comments_before"] >= 1
    assert out["existing_comments_after"] > out["existing_comments_before"]
    assert out["validation"]["existing_comments_preserved"] is True
    assert out["validation"]["new_comments_added_count"] >= 1
    assert out["validation"]["new_ai_reviewer_comments_added_count"] >= 1
    assert out["validation"]["silent_noop_suspected"] is False
    assert out["validation"]["meaningful_new_review_state"] is True
    assert out["preannotated_docx_policy"]["preserve_existing_comments"] is True
    assert out["preannotated_docx_policy"]["layer_new_comments_on_top"] is True
    assert Path(out["style_clarity_manifest_path"]).exists()
    assert Path(out["comment_response_manifest_path"]).exists()
    assert Path(out["comment_response_manifest_markdown_path"]).exists()
    assert out["comment_response_manifest"]["summary"]["existing_comments_detected"] >= 1


def test_followup_suggested_changes_append_marker_on_preannotated_docx(tmp_path: Path):
    source = tmp_path / "suggested_source.docx"
    d = Document()
    d.add_heading("Results", level=1)
    d.add_paragraph("This result sentence is too broad.\n[Suggested change] This result sentence is narrower.")
    d.save(str(source))

    out = tmp_path / "suggested_followup.docx"
    result = create_suggested_changes_docx(
        source_docx=source,
        output_docx=out,
        changes=[
            {
                "paragraph_index": 1,
                "revised_text": "This result sentence is narrowed to the tested condition.",
                "status": "applied",
            }
        ],
    )
    ins_count, del_count = _count_track_changes(out)
    assert ins_count >= 1
    assert del_count >= 1
    validation = validate_suggested_changes_docx(source, out)
    assert validation["new_tracked_change_elements_added"] >= 1
    assert validation["meaningful_new_review_state"] is True


def test_noop_regression_detected_for_preannotated_docx_with_no_new_comments(tmp_path: Path):
    source = tmp_path / "annotated.docx"
    d = Document()
    d.add_heading("Introduction", level=1)
    d.add_paragraph("This paragraph already has a reviewer note.")
    _add_existing_comment(d, 1, "Existing editor note.")
    d.save(str(source))

    reviewed = tmp_path / "reviewed.docx"
    reloaded = Document(str(source))
    reloaded.save(str(reviewed))
    validation = validate_commented_docx(source, reviewed)
    assert validation["input_preannotated"] is True
    assert validation["new_ai_reviewer_comments_added_count"] == 0
    assert validation["silent_noop_suspected"] is True
    assert validation["meaningful_new_review_state"] is False


def test_source_mode_detects_prior_ai_reviewer_suggested_docx(tmp_path: Path):
    source = tmp_path / "prior_ai.docx"
    d = Document()
    d.add_heading("Results", level=1)
    d.add_paragraph("This result is broad.\n[Suggested change] This result is limited to the tested case.")
    d.save(str(source))
    detected = detect_source_mode(source)
    assert detected["mode"] == "original_docx"
    assert detected["annotation_state"] == "prior_ai_reviewer_annotated_docx"


def test_normalize_review_artifact_text_removes_visible_suggested_blocks_only():
    raw = "Sentence one.\n[Suggested change] Tighten sentence one.\n[Suggested change - follow-up] Narrow the claim.\nSentence two."
    normalized = normalize_review_artifact_text(raw)
    assert "[Suggested change]" not in normalized
    assert "[Suggested change - follow-up]" not in normalized
    assert "Sentence one." in normalized
    assert "Sentence two." in normalized


def test_re_review_of_prior_suggested_docx_adds_followup_block_without_breaking(tmp_path: Path):
    source = tmp_path / "prior_suggested.docx"
    d = Document()
    d.add_heading("Discussion", level=1)
    d.add_paragraph(
        "This interpretation clearly demonstrates broad applicability across the reaction space.\n"
        "[Suggested change] This interpretation is limited to the tested examples."
    )
    d.save(str(source))
    parsed = parse_file(source)

    class _Item:
        def __init__(self, action: str, priority: str = "medium"):
            self.action = action
            self.priority = priority

    class _Sec:
        def __init__(self, section: str, comment: str, severity: str = "medium"):
            self.section = section
            self.comment = comment
            self.severity = severity

    class _R:
        section_specific_comments = [_Sec("Discussion", "Keep the interpretation tied to the tested evidence.")]
        extracted_action_items = [_Item("Tighten the interpretation sentence and keep the evidence boundary explicit.")]
        detailed_reviewer_comments = []
        methodological_concerns = []
        novelty_concerns = []
        citation_reference_concerns = []
        writing_organization_concerns = []
        figure_table_concerns = []
        major_weaknesses = []

    out = build_annotated_manuscript_output(source, parsed, _R(), tmp_path, run_id="reround")
    assert out["source_mode"]["annotation_state"] == "prior_ai_reviewer_annotated_docx"
    assert out["validation"]["meaningful_new_review_state"] is True
    assert out["suggested_changes_validation"]["meaningful_new_review_state"] is True
    ins_count, del_count = _count_track_changes(Path(out["suggested_changes_docx"]))
    assert ins_count >= 1
    assert del_count >= 1


def test_pre_suggested_docx_without_safe_new_rewrite_is_flagged_as_noop_risk(tmp_path: Path):
    source = tmp_path / "prior_suggested_short.docx"
    d = Document()
    d.add_heading("Discussion", level=1)
    d.add_paragraph("This interpretation is broad.\n[Suggested change] This interpretation is limited to the tested examples.")
    d.save(str(source))
    parsed = parse_file(source)

    class _Item:
        def __init__(self, action: str, priority: str = "medium"):
            self.action = action
            self.priority = priority

    class _Sec:
        def __init__(self, section: str, comment: str, severity: str = "medium"):
            self.section = section
            self.comment = comment
            self.severity = severity

    class _R:
        section_specific_comments = [_Sec("Discussion", "Keep the interpretation tied to the tested evidence.")]
        extracted_action_items = [_Item("Tighten the interpretation sentence and keep the evidence boundary explicit.")]
        detailed_reviewer_comments = []
        methodological_concerns = []
        novelty_concerns = []
        citation_reference_concerns = []
        writing_organization_concerns = []
        figure_table_concerns = []
        major_weaknesses = []

    out = build_annotated_manuscript_output(source, parsed, _R(), tmp_path, run_id="reround_short")
    assert out["source_mode"]["annotation_state"] == "prior_ai_reviewer_annotated_docx"
    assert out["validation"]["meaningful_new_review_state"] is True
    assert out["suggested_changes_validation"]["input_preannotated"] is True
    assert out["suggested_changes_validation"]["silent_noop_suspected"] is True
    assert out["suggested_changes_validation"]["meaningful_new_review_state"] is False
