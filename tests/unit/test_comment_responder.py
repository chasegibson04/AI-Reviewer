from __future__ import annotations

import json
from pathlib import Path

from docx import Document

from ai_reviewer.models.base import ChatRequest
from ai_reviewer.review.comment_responder import (
    RESPONSE_NEEDS_CLARIFICATION,
    RESPONSE_TEXT_FIX,
    render_comment_response_manifest_markdown,
    run_existing_comment_responder,
)
from ai_reviewer.tools.docx_tools import extract_existing_docx_comments


class _MockProvider:
    def __init__(self, responses: list[str], models: list[str] | None = None):
        self._responses = list(responses)
        self._models = models or ["gemma4:31b"]

    def list_models(self) -> list[str]:
        return list(self._models)

    def chat(self, request: ChatRequest):
        if not self._responses:
            raise RuntimeError("No response queued")

        class _Resp:
            def __init__(self, content: str):
                self.content = content

        return _Resp(self._responses.pop(0))


class _AlwaysFailProvider:
    def __init__(self, models: list[str] | None = None):
        self._models = models or ["gemma4:31b"]
        self.calls = 0

    def list_models(self) -> list[str]:
        return list(self._models)

    def chat(self, request: ChatRequest):
        self.calls += 1
        raise RuntimeError("empty_response")


def _write_docx_with_comment(path: Path, paragraph_text: str, comment_text: str) -> None:
    doc = Document()
    p = doc.add_paragraph(paragraph_text)
    if not p.runs:
        p.add_run(paragraph_text)
    comment = doc.comments.add_comment(text=comment_text, author="Editor", initials="ED")
    p.runs[0].mark_comment_range(p.runs[-1], comment.comment_id)
    doc.save(str(path))


def test_extract_existing_docx_comments_returns_anchor_and_location(tmp_path: Path):
    source = tmp_path / "annotated.docx"
    _write_docx_with_comment(
        source,
        "This sentence is overly long and should be split for readability in the methods description.",
        "Please split this sentence for clarity.",
    )
    rows = extract_existing_docx_comments(source)
    assert len(rows) >= 1
    assert rows[0]["comment_text"]
    assert rows[0]["anchor_text"]
    assert isinstance(rows[0]["paragraph_index"], int)


def test_existing_comment_responder_generates_manifest_entries(tmp_path: Path):
    source = tmp_path / "annotated.docx"
    paragraph = "This sentence is overly long and should be split for readability in the methods description."
    _write_docx_with_comment(source, paragraph, "Please split this sentence for clarity.")
    manifest = run_existing_comment_responder(
        source_docx=source,
        paragraphs=[paragraph],
        section_by_idx={0: "methods"},
        provider=None,
        model=None,
        timeout_seconds=20,
    )
    assert manifest["summary"]["existing_comments_detected"] >= 1
    assert manifest["summary"]["responses_generated"] >= 1
    first = manifest["responses"][0]
    assert first["response_type"] in {RESPONSE_TEXT_FIX, "response_strategy", RESPONSE_NEEDS_CLARIFICATION, "already_addressed"}
    assert first["comment_text"]
    assert first["proposed_response"]


def test_existing_comment_responder_uses_llm_and_validates_text_fix(tmp_path: Path):
    source = tmp_path / "annotated.docx"
    paragraph = "This sentence is overly long and should be split for readability in the methods description."
    _write_docx_with_comment(source, paragraph, "Please split this sentence for clarity.")
    provider = _MockProvider(
        responses=[
            json.dumps(
                {
                    "response_type": "text_fix",
                    "proposed_response": "Apply a line edit that separates setup from outcome.",
                    "proposed_rewrite": "This sentence is overly long. Split setup from outcome for readability.",
                    "confidence": "high",
                }
            )
        ]
    )
    manifest = run_existing_comment_responder(
        source_docx=source,
        paragraphs=[paragraph],
        section_by_idx={0: "methods"},
        provider=provider,
        model="fallback-model",
        timeout_seconds=20,
    )
    assert manifest["model"]["selected_model"] == "gemma4:31b"
    assert manifest["summary"]["responses_generated"] >= 1
    assert manifest["responses"][0]["response_type"] in {RESPONSE_TEXT_FIX, "response_strategy"}


def test_comment_response_manifest_markdown_is_rendered():
    manifest = {
        "summary": {"existing_comments_detected": 1, "responses_generated": 1},
        "responses": [
            {
                "comment_id": "1",
                "section": "methods",
                "response_type": "response_strategy",
                "confidence": "medium",
                "proposed_response": "Clarify parameter values in the same sentence.",
                "proposed_rewrite": None,
            }
        ],
    }
    text = render_comment_response_manifest_markdown(manifest)
    assert "# Existing Comment Response Manifest" in text
    assert "| Comment ID | Section | Response Type | Confidence | Proposed Response |" in text


def test_existing_comment_responder_fail_fast_disables_llm_after_first_error(tmp_path: Path):
    source = tmp_path / "annotated.docx"
    paragraph = "This sentence is overly long and should be split for readability in the methods description."
    doc = Document()
    p1 = doc.add_paragraph(paragraph)
    p1.add_run("")
    c1 = doc.comments.add_comment(text="Please split this sentence for clarity.", author="Editor", initials="ED")
    p1.runs[0].mark_comment_range(p1.runs[-1], c1.comment_id)
    p2 = doc.add_paragraph(paragraph)
    p2.add_run("")
    c2 = doc.comments.add_comment(text="Please improve the flow in this sentence.", author="Editor", initials="ED")
    p2.runs[0].mark_comment_range(p2.runs[-1], c2.comment_id)
    doc.save(str(source))
    provider = _AlwaysFailProvider()
    manifest = run_existing_comment_responder(
        source_docx=source,
        paragraphs=[paragraph, paragraph],
        section_by_idx={0: "methods", 1: "methods"},
        provider=provider,
        model="fallback-model",
        timeout_seconds=20,
    )
    assert provider.calls == 1
    assert manifest["model"]["fail_fast_fallback_triggered"] is True


def test_existing_comment_responder_demotes_overlong_text_fix_rewrite(tmp_path: Path):
    source = tmp_path / "annotated.docx"
    paragraph = (
        "Chemical synthesis is a primary bottleneck in drug development. "
        "High-throughput experimentation is a widely practiced method for discovery and optimization."
    )
    _write_docx_with_comment(source, paragraph, "Please split this sentence for clarity.")
    long_rewrite = (
        "Chemical synthesis is a primary bottleneck in drug development. "
        "High-throughput experimentation is a widely practiced method for discovery and optimization. "
        "This additional sentence makes the rewrite too long for a local text-fix response."
    )
    provider = _MockProvider(
        responses=[
            json.dumps(
                {
                    "response_type": "text_fix",
                    "proposed_response": "Apply a local rewrite.",
                    "proposed_rewrite": long_rewrite,
                    "confidence": "high",
                }
            )
        ]
    )
    manifest = run_existing_comment_responder(
        source_docx=source,
        paragraphs=[paragraph],
        section_by_idx={0: "introduction"},
        provider=provider,
        model="fallback-model",
        timeout_seconds=20,
    )
    row = manifest["responses"][0]
    assert row["response_type"] == "response_strategy"
    assert row["proposed_rewrite"] is None


def test_existing_comment_responder_never_emits_literal_none_rewrite(tmp_path: Path):
    source = tmp_path / "annotated.docx"
    paragraph = "Prior work reported this trend in similar systems."
    _write_docx_with_comment(source, paragraph, "Need citation.")
    manifest = run_existing_comment_responder(
        source_docx=source,
        paragraphs=[paragraph],
        section_by_idx={0: "introduction"},
        provider=None,
        model=None,
        timeout_seconds=20,
    )
    row = manifest["responses"][0]
    assert row["proposed_rewrite"] is None or row["proposed_rewrite"] != "None"
